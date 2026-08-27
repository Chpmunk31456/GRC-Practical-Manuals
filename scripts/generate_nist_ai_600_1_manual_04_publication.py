#!/usr/bin/env python3
"""Generate trilingual Manual 04 publication-QA candidates using the proven Manual 03 engine.

This adapter deliberately reuses document rendering/conversion primitives while overriding
Manual 04 source discovery, heading grammar, metadata, assurance boundaries, and reports.
Generation is not human semantic approval or release approval.
"""
from __future__ import annotations

import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

import generate_nist_ai_rmf_manual_03_publication as core

ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "01-foundations" / "NIST_AI_600-1_Generative_AI_Profile"
BASELINE = ROOT / ".compliance" / "nist-ai-600-1-manual-04-baseline.json"

LANG_META = {
    "en": {
        "title": "Manual 04 - NIST AI 600-1 Generative AI Profile Implementation",
        "word_lang": "en-US",
        "filename": "Manual_04_NIST_AI_600-1_Generative_AI_Profile_EN",
        "status": "CONTROLLED PUBLICATION QA CANDIDATE",
    },
    "es-419": {
        "title": "Manual 04 - Implementación del Perfil de IA Generativa NIST AI 600-1",
        "word_lang": "es-419",
        "filename": "Manual_04_NIST_AI_600-1_Generative_AI_Profile_ES-419",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACION",
    },
    "pt-BR": {
        "title": "Manual 04 - Implementação do Perfil de IA Generativa NIST AI 600-1",
        "word_lang": "pt-BR",
        "filename": "Manual_04_NIST_AI_600-1_Generative_AI_Profile_PT-BR",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACAO",
    },
}

CHAPTER_RE = re.compile(r"(?mi)^##\s+(?:Chapter|Cap[ií]tulo)\s+([0-9]{1,2})\s+[—-]\s+(.+?)\s*$")


def split_chapters(text: str) -> dict[int, str]:
    matches = list(CHAPTER_RE.finditer(text))
    chapters: dict[int, str] = {}
    for index, match in enumerate(matches):
        number = int(match.group(1))
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        chapters[number] = text[match.start():end].rstrip() + "\n"
    return chapters


def chapter_dir(language: str) -> Path:
    return MANUAL / ("English" if language == "en" else language) / "source"


def find_localized_chapters(language: str) -> tuple[str, list[str]]:
    paths = sorted(chapter_dir(language).glob("*.md"))
    chapters: dict[int, str] = {}
    used: list[str] = []
    for path in paths:
        text = path.read_text(encoding="utf-8")
        found = split_chapters(text)
        if not found:
            continue
        for number, body in found.items():
            if number in chapters and chapters[number] != body:
                raise ValueError(f"conflicting chapter {number} in {language}: {path}")
            chapters[number] = body
        used.append(str(path.relative_to(ROOT)))
    expected = set(range(1, 33))
    missing = sorted(expected - set(chapters))
    extra = sorted(set(chapters) - expected)
    if missing or extra:
        raise ValueError(f"{language} chapter inventory invalid; missing={missing}, extra={extra}")
    return "\n".join(chapters[n].rstrip() for n in range(1, 33)) + "\n", used


def find_implementation(language: str) -> tuple[str, str]:
    if language == "en":
        path = MANUAL / "MANUAL_04_IMPLEMENTATION_PATHS.md"
    else:
        base = chapter_dir(language)
        names = {
            "es-419": "RUTAS_DE_IMPLEMENTACION_MANUAL_04.md",
            "pt-BR": "CAMINHOS_DE_IMPLEMENTACAO_MANUAL_04.md",
        }
        path = base / names[language]
    if not path.is_file():
        raise FileNotFoundError(path)
    text = path.read_text(encoding="utf-8")
    if len(text) < 3000:
        raise ValueError(f"implementation entry unexpectedly small for {language}: {path}")
    return text, str(path.relative_to(ROOT))


def add_footer(section, language: str):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual 04 | Controlled publication QA candidate | ")
    core.set_run_font(run, language, size=8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def build_docx(source: core.EditionSource, out_path: Path, image_dir: Path, source_head: str):
    language = source.language
    meta = LANG_META[language]
    doc = Document()
    core.set_document_defaults(doc, language)
    add_footer(doc.sections[0], language)

    props = doc.core_properties
    props.title = meta["title"]
    props.subject = f"Manual 04 controlled publication-QA artifact; language {language}; source head {source_head}"
    props.author = "Alberto (Al) Leiva"
    props.keywords = "NIST AI 600-1, Generative AI Profile, AI RMF, GOVERN, MAP, MEASURE, MANAGE, GRC"
    props.comments = "Generation and QA do not establish certification, legal compliance, trustworthy-AI achievement, or an audit opinion."

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["title"])
    core.set_run_font(run, language, size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta["status"])
    core.set_run_font(run, language, size=10, bold=True)
    control = doc.add_paragraph()
    control.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = control.add_run(
        f"Controlled source revision: {source_head} | Language: {language} | "
        "NIST AI 600-1 Generative AI Profile with AI RMF 1.0 parent framework; version-aware and source-watched."
    )
    core.set_run_font(run, language, size=8.5)

    boundary = doc.add_paragraph()
    boundary.paragraph_format.left_indent = Inches(0.2)
    boundary.paragraph_format.right_indent = Inches(0.2)
    run = boundary.add_run(
        "Assurance boundary: This voluntary implementation manual does not certify a GAI system, "
        "establish legal compliance, prove trustworthy-AI achievement, or provide an audit opinion; "
        "not every suggested action applies to every actor, system, use case, or lifecycle context."
    )
    core.set_run_font(run, language, size=9, bold=True)

    graphic_counter = [0]
    h = doc.add_paragraph("Implementation paths and operating model", style="Heading 1")
    core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.implementation_text, language, image_dir, graphic_counter, "Implementation")
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1], language)
    h = doc.add_paragraph("Controlled 32-chapter manual", style="Heading 1")
    core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.chapter_text, language, image_dir, graphic_counter, "Chapter")

    required_graphics = int(json.loads(BASELINE.read_text(encoding="utf-8"))["required_visuals_in_implementation_entry"])
    if graphic_counter[0] < required_graphics:
        raise ValueError(f"{language} has too few memory graphics: {graphic_counter[0]} < {required_graphics}")

    for p in doc.paragraphs:
        for r in p.runs:
            core.set_run_font(r, language)
    doc.save(out_path)
    return graphic_counter[0]


def inspect_docx(path: Path, language: str, expected_graphics: int) -> dict:
    import zipfile
    if not zipfile.is_zipfile(path):
        raise ValueError(f"not a valid DOCX ZIP: {path}")
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    chapter_numbers = [int(x) for x in re.findall(r"(?mi)^(?:Chapter|Cap[ií]tulo)\s+([0-9]{1,2})\s+[—-]", text)]
    chapter_set = sorted(set(n for n in chapter_numbers if 1 <= n <= 32))
    if len(text) < 30_000:
        raise ValueError(f"DOCX text unexpectedly small: {path}")
    if chapter_set != list(range(1, 33)):
        raise ValueError(f"DOCX chapter inventory incomplete for {language}: {chapter_set}")
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
        styles = zf.read("word/styles.xml").decode("utf-8", errors="replace")
    alt_count = len(re.findall(r"\bdescr=\"[^\"]+\"", xml))
    if alt_count < expected_graphics:
        raise ValueError(f"DOCX image alt-text count {alt_count} < graphics {expected_graphics}: {path}")
    if LANG_META[language]["word_lang"] not in styles and LANG_META[language]["word_lang"] not in xml:
        raise ValueError(f"DOCX language metadata missing for {language}: {path}")
    return {
        "file": path.name, "bytes": path.stat().st_size, "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables), "inline_shapes": len(doc.inline_shapes), "alt_text_entries": alt_count,
        "chapter_count": len(chapter_set), "sha256": core.sha256(path), "status": "PASS",
    }


def write_reports(output_root: Path, sources: dict[str, core.EditionSource], file_results: dict, page_rows: list[dict], source_head: str):
    qa_dir = output_root / "qa"
    qa_dir.mkdir(parents=True, exist_ok=True)
    manifest = {
        "schema_version": "1.0",
        "manual": "Manual 04 - NIST AI 600-1 Generative AI Profile Implementation",
        "source_head": source_head,
        "source_branch": "build/nist-ai-600-1-manual-04-2026",
        "controlled_baseline": "NIST AI 600-1 Generative AI Profile with AI RMF 1.0 parent framework",
        "assurance_boundary": "Publication QA does not establish legal compliance, certification, trustworthy-AI achievement, universal applicability, or an audit opinion.",
        "editions": {},
    }
    for language, source in sources.items():
        manifest["editions"][language] = {
            "chapter_sources": source.chapter_files,
            "implementation_source": source.implementation_file,
            "chapter_count": 32,
            "mermaid_source_blocks": len(core.MERMAID_RE.findall(source.chapter_text)) + len(core.MERMAID_RE.findall(source.implementation_text)),
            "artifacts": file_results[language],
        }
    (qa_dir / "MANUAL_04_PUBLICATION_REPORT.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    (qa_dir / "MANUAL_04_PUBLICATION_REPORT.md").write_text(
        "# Manual 04 Publication Processing Report\n\n"
        f"- Controlled source head: `{source_head}`\n"
        "- Candidate languages: English (`en`), Latin American Spanish (`es-419`), Brazilian Portuguese (`pt-BR`)\n"
        "- Automated document-processing status: **PASS**\n"
        "- Release status: **QA CANDIDATE - human semantic/accessibility/release controls remain authoritative**\n\n"
        "## Automated checks\n\n"
        "- 32 chapters detected in each DOCX candidate.\n"
        "- DOCX packages, language metadata, and image alternative text validated.\n"
        "- Every PDF page contains extractable text and rendered successfully.\n"
        "- SHA-256 provenance recorded for all six publication candidates.\n"
        "- Complete page renders and contact sheets generated for human visual review.\n\n"
        "## Assurance boundary\n\n"
        "Successful conversion and automated QA do not constitute human semantic approval, legal compliance, certification, universal applicability, trustworthy-AI achievement, or an audit opinion.\n",
        encoding="utf-8",
    )
    with (qa_dir / "MANUAL_04_PAGE_QA.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["pdf", "page", "width_pt", "height_pt", "text_chars", "render", "automated_status"])
        writer.writeheader(); writer.writerows(page_rows)
    candidates = sorted((output_root / "publication").glob("*.docx")) + sorted((output_root / "publication").glob("*.pdf"))
    with (qa_dir / "MANUAL_04_SHA256SUMS.txt").open("w", encoding="utf-8") as handle:
        for path in candidates:
            handle.write(f"{core.sha256(path)}  {path.name}\n")


core.MANUAL = MANUAL
core.BASELINE = BASELINE
core.LANG_META = LANG_META
core.CHAPTER_RE = CHAPTER_RE
core.split_chapters = split_chapters
core.find_localized_chapters = find_localized_chapters
core.find_implementation = find_implementation
core.add_footer = add_footer
core.build_docx = build_docx
core.inspect_docx = inspect_docx
core.write_reports = write_reports

if __name__ == "__main__":
    raise SystemExit(core.main())
