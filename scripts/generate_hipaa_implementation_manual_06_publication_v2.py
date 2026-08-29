#!/usr/bin/env python3
"""Manual 06 publication generator v2.

Wraps the established HIPAA generator to localize generator-owned control text,
figure captions, and localized document metadata before human review. Controlled
source meaning and current-law/proposed-rule boundaries are unchanged.
"""
from __future__ import annotations

from docx import Document

import generate_hipaa_implementation_manual_06_publication as v1

_original_build_docx = v1.build_docx

CONTROL = {
    "en": {
        "prefix": "Controlled source revision:",
        "tail": "current-law/proposed-rule source-state controls retained.",
        "caption": "HIPAA implementation memory graphic",
    },
    "es-419": {
        "prefix": "Revisión de fuente controlada:",
        "tail": "se conservan los controles de estado de fuentes de normativa vigente y regla propuesta.",
        "caption": "Gráfico de memoria de implementación de HIPAA",
    },
    "pt-BR": {
        "prefix": "Revisão da fonte controlada:",
        "tail": "são mantidos os controles de estado de fontes da norma vigente e da regra proposta.",
        "caption": "Gráfico de memória de implementação da HIPAA",
    },
}


def replace_paragraph_text(paragraph, text: str, language: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
    else:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    for run in paragraph.runs:
        v1.core.set_run_font(run, language)


def build_docx(source, out_path, image_dir, source_head: str):
    count = _original_build_docx(source, out_path, image_dir, source_head)
    language = source.language
    meta = v1.LANG_META[language]
    loc = CONTROL[language]
    doc = Document(out_path)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Controlled source revision:"):
            replace_paragraph_text(
                paragraph,
                f"{loc['prefix']} {source_head} | {meta['language_label']} {language} | {loc['tail']}",
                language,
            )
            continue
        if language in {"es-419", "pt-BR"} and text.startswith("Figure ") and "Implementation memory graphic" in text:
            number = text.split(".", 1)[0].split()[-1]
            replace_paragraph_text(paragraph, f"Figura {number}. {loc['caption']}", language)

    if language == "es-419":
        doc.core_properties.subject = f"Artefacto controlado de QA de publicación del Manual 06; idioma es-419; revisión de fuente {source_head}"
        doc.core_properties.comments = "La generación automatizada no establece condición jurídica, cumplimiento de HIPAA, notificabilidad de una brecha ni una opinión de auditoría."
    elif language == "pt-BR":
        doc.core_properties.subject = f"Artefato controlado de QA de publicação do Manual 06; idioma pt-BR; revisão da fonte {source_head}"
        doc.core_properties.comments = "A geração automatizada não estabelece status jurídico, conformidade HIPAA, notificabilidade de uma violação nem opinião de auditoria."

    doc.save(out_path)
    return count


v1.build_docx = build_docx
v1.core.build_docx = build_docx

if __name__ == "__main__":
    raise SystemExit(v1.core.main())
