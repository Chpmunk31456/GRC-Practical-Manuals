#!/usr/bin/env python3
"""Build the exact Manual 47 trilingual DOCX/PDF publication candidate."""

from pathlib import Path
from docx import Document
from docx.shared import Pt
import hashlib
import json
import os
import re
import subprocess

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "10-ai-governance/Manual_47_EU_AI_Act_Training"
OUT = ROOT / "qa/manual47-publication-candidate"
OUT.mkdir(parents=True, exist_ok=True)

EN_COMPONENTS = [
    BASE / "README.md",
    BASE / "TRAINING_MODULES.md",
    BASE / "PRACTICAL_SCENARIOS.md",
    BASE / "CONTROL_EVIDENCE_WORKBOOK.md",
]
LOCALE_SOURCES = {
    "es-419": BASE / "controlled/es-419/MANUAL_47_CONTROLLED_ES-419.md",
    "pt-BR": BASE / "controlled/pt-BR/MANUAL_47_CONTROLLED_PT-BR.md",
}
NAMES = {
    "en": "Manual_47_EU_AI_Act_Training_Operationalization_EN",
    "es-419": "Manual_47_EU_AI_Act_Training_Operationalization_ES-419",
    "pt-BR": "Manual_47_EU_AI_Act_Training_Operationalization_PT-BR",
}


def read_required(path: Path) -> str:
    if not path.is_file():
        raise SystemExit(f"required controlled source missing: {path.relative_to(ROOT)}")
    text = path.read_text(encoding="utf-8")
    if len(text.strip()) < 700:
        raise SystemExit(f"controlled source is unexpectedly short: {path.relative_to(ROOT)}")
    return text


def english_source() -> str:
    text = "\n\n---\n\n".join(read_required(path) for path in EN_COMPONENTS)
    required = [
        "EU AI Act",
        "Training Modules",
        "Practical Scenarios",
        "Control and Evidence Workbook",
        "prohibited",
        "high-risk",
        "general-purpose",
        "human oversight",
    ]
    lower = text.lower()
    for marker in required:
        if marker.lower() not in lower:
            raise SystemExit(f"English controlled package missing marker: {marker}")
    return text


def add_markdown(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            doc.add_heading(heading.group(2).strip(), level=min(len(heading.group(1)), 3))
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        elif line != "---":
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)


sources = {
    "en": english_source(),
    "es-419": read_required(LOCALE_SOURCES["es-419"]),
    "pt-BR": read_required(LOCALE_SOURCES["pt-BR"]),
}

manifest = {
    "manual": 47,
    "title": "EU AI Act Training & Operationalization",
    "source_commit": os.environ.get("GITHUB_SHA", ""),
    "artifacts": [],
}

for locale, text in sources.items():
    doc = Document()
    doc.styles["Normal"].font.name = "Liberation Sans"
    doc.styles["Normal"].font.size = Pt(10.5)
    add_markdown(doc, text)
    stem = NAMES[locale]
    docx = OUT / f"{stem}.docx"
    doc.save(docx)
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx)],
        check=True,
    )
    pdf = OUT / f"{stem}.pdf"
    for path in (docx, pdf):
        if not path.is_file() or path.stat().st_size == 0:
            raise SystemExit(f"empty or missing artifact: {path}")
        data = path.read_bytes()
        manifest["artifacts"].append({
            "locale": locale,
            "file": path.name,
            "sha256": hashlib.sha256(data).hexdigest(),
            "bytes": len(data),
        })

(OUT / "MANUAL_47_CANDIDATE_MANIFEST.json").write_text(
    json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
)
print(json.dumps(manifest, indent=2, ensure_ascii=False))
