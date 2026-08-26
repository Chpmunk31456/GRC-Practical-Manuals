#!/usr/bin/env python3
"""Generate controlled publication-QA candidates for Manual 03.

The controlled Markdown sources remain authoritative. This script creates
DOCX/PDF candidates and machine-readable QA evidence in an isolated output
folder. Generation is not release approval and does not establish legal
compliance, certification, trustworthy-AI achievement, or an audit opinion.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "01-foundations" / "NIST_AI_RMF_1.0"
BASELINE = ROOT / ".compliance" / "nist-ai-rmf-manual-03-baseline.json"

LANG_META = {
    "en": {
        "title": "Manual 03 - NIST AI Risk Management Framework Implementation",
        "word_lang": "en-US",
        "filename": "Manual_03_NIST_AI_RMF_Implementation_EN",
        "status": "CONTROLLED PUBLICATION QA CANDIDATE",
    },
    "es-419": {
        "title": "Manual 03 - Implementacion del Marco de Gestion de Riesgos de IA de NIST",
        "word_lang": "es-419",
        "filename": "Manual_03_NIST_AI_RMF_Implementation_ES-419",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACION",
    },
    "pt-BR": {
        "title": "Manual 03 - Implementacao do NIST AI Risk Management Framework",
        "word_lang": "pt-BR",
        "filename": "Manual_03_NIST_AI_RMF_Implementation_PT-BR",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACAO",
    },
}

CHAPTER_RE = re.compile(r"(?m)^#\s+([0-9]{1,2})\.\s+(.+?)\s*$")
MERMAID_RE = re.compile(r"(?ms)^```mermaid\s*\n(.*?)^```\s*$")
LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
TABLE_DIVIDER_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*$")


@dataclass
class EditionSource:
    language: str
    chapter_text: str
    implementation_text: str
    chapter_files: list[str]
    implementation_file: str


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def clean_inline(text: str) -> str:
    text = html.unescape(text.strip())
    text = LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)
    text = BOLD_RE.sub(r"\1", text)
    text = INLINE_CODE_RE.sub(r"\1", text)
    text = text.replace("<br>", " ").replace("<br/>", " ").replace("<br />", " ")
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\s+", " ", text).strip()


def split_chapters(text: str) -> dict[int, str]:
    matches = list(CHAPTER_RE.finditer(text))
    chapters: dict[int, str] = {}
    for index, match in enumerate(matches):
        number = int(match.group(1))
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        chapters[number] = text[match.start():end].rstrip() + "\n"
    return chapters


def find_localized_chapters(language: str) -> tuple[str, list[str]]:
    if language == "en":
        paths = [MANUAL / p for p in json.loads(BASELINE.read_text(encoding="utf-8"))["english_source_parts"]]
    else:
        candidates = []
        for base in [MANUAL / "translations" / language / "source", MANUAL / language / "source"]:
            if base.is_dir():
                candidates.extend(sorted(base.glob("*.md")))
        paths = [p for p in candidates if split_chapters(p.read_text(encoding="utf-8"))]

    chapters: dict[int, str] = {}
    used: list[str] = []
    for path in paths:
        contents = path.read_text(encoding="utf-8")
        found = split_chapters(contents)
        if not found:
            continue
        for number, body in found.items():
            if number in chapters and chapters[number] != body:
                raise ValueError(f"conflicting chapter {number} in {language}: {path}")
            chapters[number] = body
        used.append(str(path.relative_to(ROOT)))

    expected = set(range(1, 33))
    missing = sorted(expected - set(chapters))
    extra = sorted(set(chapters) - expected)
    if missing or extra:
        raise ValueError(f"{language} chapter inventory invalid; missing={missing}, extra={extra}")
    return "\n".join(chapters[n].rstrip() for n in range(1, 33)) + "\n", used


def find_implementation(language: str) -> tuple[str, str]:
    if language == "en":
        path = MANUAL / "MANUAL_03_IMPLEMENTATION_PATHS.md"
    else:
        base = MANUAL / "translations" / language / "source"
        patterns = {
            "es-419": ["RUTAS_DE_IMPLEMENTACION_MANUAL_03.md", "*IMPLEMENTACION*03*.md"],
            "pt-BR": ["CAMINHOS_DE_IMPLEMENTACAO_MANUAL_03.md", "*IMPLEMENTACAO*03*.md"],
        }[language]
        path = None
        for pattern in patterns:
            matches = sorted(base.glob(pattern))
            if matches:
                path = matches[0]
                break
        if path is None:
            raise FileNotFoundError(f"localized implementation entry not found for {language}")
    text = path.read_text(encoding="utf-8")
    if len(text) < 3000:
        raise ValueError(f"implementation entry unexpectedly small for {language}: {path}")
    return text, str(path.relative_to(ROOT))


def load_sources() -> dict[str, EditionSource]:
    result = {}
    for language in LANG_META:
        chapter_text, chapter_files = find_localized_chapters(language)
        implementation_text, implementation_file = find_implementation(language)
        result[language] = EditionSource(
            language=language,
            chapter_text=chapter_text,
            implementation_text=implementation_text,
            chapter_files=chapter_files,
            implementation_file=implementation_file,
        )
    return result


def font(size: int, bold: bool = False):
    names = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in names:
        if Path(candidate).is_file():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def parse_mermaid(block: str) -> tuple[list[tuple[str, str]], list[tuple[str, str]]]:
    node_labels: dict[str, str] = {}
    order: list[str] = []
    edges: list[tuple[str, str]] = []

    node_pattern = re.compile(r"\b([A-Za-z][A-Za-z0-9_]*)\s*(?:\[\[|\[|\(|\{)\s*[\"']?(.+?)[\"']?\s*(?:\]\]|\]|\)|\})")
    edge_pattern = re.compile(r"\b([A-Za-z][A-Za-z0-9_]*)\b\s*(?:-->|---|-.->|==>)\s*(?:\|[^|]*\|\s*)?\b([A-Za-z][A-Za-z0-9_]*)\b")

    for raw in block.splitlines():
        line = raw.strip()
        if not line or line.startswith("flowchart") or line.startswith("graph") or line.startswith("%%"):
            continue
        for match in node_pattern.finditer(line):
            key = match.group(1)
            label = re.sub(r"<br\s*/?>", " ", match.group(2), flags=re.I)
            label = label.strip('"\' ')
            if key not in node_labels:
                order.append(key)
            node_labels[key] = clean_inline(label)[:150]
        for match in edge_pattern.finditer(line):
            a, b = match.group(1), match.group(2)
            edges.append((a, b))
            for key in (a, b):
                if key not in node_labels:
                    node_labels[key] = key
                    order.append(key)

    if not order:
        order = ["FLOW"]
        node_labels["FLOW"] = "AI risk-management relationship"
    return [(key, node_labels[key]) for key in order], edges


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    line = words[0]
    for word in words[1:]:
        trial = f"{line} {word}"
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            line = trial
        else:
            lines.append(line)
            line = word
    lines.append(line)
    return lines[:5]


def render_mermaid_memory_graphic(block: str, out_path: Path, title: str) -> str:
    nodes, edges = parse_mermaid(block)
    width = 1500
    box_width = 1080
    box_height = 145
    gap = 70
    margin_top = 150
    height = margin_top + len(nodes) * (box_height + gap) + 70
    canvas = Image.new("RGB", (width, max(height, 550)), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = font(34, bold=True)
    body_font = font(28)
    small_font = font(22)

    draw.text((70, 45), title[:90], fill="black", font=title_font)
    x = (width - box_width) // 2
    positions: dict[str, tuple[int, int, int, int]] = {}
    for idx, (key, label) in enumerate(nodes):
        y = margin_top + idx * (box_height + gap)
        rect = (x, y, x + box_width, y + box_height)
        positions[key] = rect
        draw.rounded_rectangle(rect, radius=18, outline="black", width=3, fill="#f6f6f6")
        lines = wrap_text(draw, label, body_font, box_width - 80)
        line_h = 34
        total = line_h * len(lines)
        ty = y + (box_height - total) // 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=body_font)
            tx = x + (box_width - (bbox[2] - bbox[0])) // 2
            draw.text((tx, ty), line, fill="black", font=body_font)
            ty += line_h
        draw.text((x + 16, y + 10), key, fill="#555555", font=small_font)

    for a, b in edges:
        if a not in positions or b not in positions:
            continue
        ra, rb = positions[a], positions[b]
        start = ((ra[0] + ra[2]) // 2, ra[3])
        end = ((rb[0] + rb[2]) // 2, rb[1])
        if end[1] <= start[1]:
            continue
        draw.line((start, end), fill="black", width=4)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 12, ey - 20), (ex + 12, ey - 20)], fill="black")

    canvas.save(out_path, format="PNG", optimize=True)
    return f"Memory graphic: {title}. Diagram contains {len(nodes)} labeled nodes and {len(edges)} directed relationships."


def set_run_font(run, language: str, size: float | None = None, bold: bool | None = None):
    run.font.name = "Arial"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), "Arial")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), LANG_META[language]["word_lang"])


def set_paragraph_keep(paragraph, keep_next: bool = False):
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:keepLines")) is None:
        ppr.append(OxmlElement("w:keepLines"))
    if keep_next and ppr.find(qn("w:keepNext")) is None:
        ppr.append(OxmlElement("w:keepNext"))


def set_cell_shading(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_image_alt_text(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title[:250])
    doc_pr.set("descr", description[:1000])


def add_hyperlink(paragraph, text: str, url: str, language: str):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    t = OxmlElement("w:t")
    t.text = text
    run_el.extend([rpr, t])
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def add_inline_text(paragraph, text: str, language: str):
    pos = 0
    for match in LINK_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(clean_inline(text[pos:match.start()]))
            set_run_font(run, language)
        add_hyperlink(paragraph, clean_inline(match.group(1)), match.group(2), language)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(clean_inline(text[pos:]))
        set_run_font(run, language)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or "|" not in lines[start] or not TABLE_DIVIDER_RE.match(lines[start + 1]):
        return None
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and "|" in lines[i] and lines[i].strip():
        if not TABLE_DIVIDER_RE.match(lines[i]):
            row = [clean_inline(x) for x in lines[i].strip().strip("|").split("|")]
            rows.append(row)
        i += 1
    return rows, i


def add_markdown(document: Document, text: str, language: str, image_dir: Path, graphic_counter: list[int], section_label: str):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue

        if stripped == "```mermaid":
            block_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                block_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            graphic_counter[0] += 1
            number = graphic_counter[0]
            png = image_dir / f"{language.replace('-', '_')}_graphic_{number:02d}.png"
            alt = render_mermaid_memory_graphic("\n".join(block_lines), png, f"Manual 03 memory graphic {number}")
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shape = p.add_run().add_picture(str(png), width=Inches(6.2))
            set_image_alt_text(shape, f"Manual 03 memory graphic {number}", alt)
            set_paragraph_keep(p, keep_next=True)
            cap = document.add_paragraph(f"Figure {number}. {section_label} memory graphic", style="Caption")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_keep(cap, keep_next=True)
            continue

        if stripped.startswith("```"):
            fence = stripped
            code_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = document.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            continue

        table = parse_table(lines, i)
        if table:
            rows, next_i = table
            max_cols = max(len(r) for r in rows) if rows else 0
            if rows and max_cols:
                tbl = document.add_table(rows=len(rows), cols=max_cols)
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r, row in enumerate(rows):
                    for c in range(max_cols):
                        cell = tbl.cell(r, c)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                        cell.text = row[c] if c < len(row) else ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                set_run_font(run, language, size=8.5, bold=(r == 0))
                        if r == 0:
                            set_cell_shading(cell, "E7E6E6")
                    if r == 0:
                        trpr = tbl.rows[r]._tr.get_or_add_trPr()
                        tbl_header = OxmlElement("w:tblHeader")
                        tbl_header.set(qn("w:val"), "true")
                        trpr.append(tbl_header)
            i = next_i
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            title = clean_inline(heading.group(2))
            p = document.add_paragraph(title, style=f"Heading {level}")
            for run in p.runs:
                set_run_font(run, language, size={1: 18, 2: 15, 3: 12, 4: 11}[level], bold=True)
            set_paragraph_keep(p, keep_next=True)
            i += 1
            continue

        if re.match(r"^[-*+]\s+", stripped):
            p = document.add_paragraph(style="List Bullet")
            add_inline_text(p, re.sub(r"^[-*+]\s+", "", stripped), language)
            set_paragraph_keep(p)
            i += 1
            continue

        if re.match(r"^\d+[.)]\s+", stripped):
            p = document.add_paragraph(style="List Number")
            add_inline_text(p, re.sub(r"^\d+[.)]\s+", "", stripped), language)
            set_paragraph_keep(p)
            i += 1
            continue

        if stripped.startswith(">"):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_inline_text(p, stripped.lstrip("> "), language)
            set_paragraph_keep(p)
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("```") or re.match(r"^[-*+]\s+", nxt) or re.match(r"^\d+[.)]\s+", nxt):
                break
            if "|" in nxt and i + 1 < len(lines) and TABLE_DIVIDER_RE.match(lines[i + 1]):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = document.add_paragraph()
        add_inline_text(p, " ".join(paragraph_lines), language)
        set_paragraph_keep(p)


def set_document_defaults(document: Document, language: str):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for level in range(1, 5):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt({1: 18, 2: 15, 3: 12, 4: 11}[level])
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    for style in document.styles:
        rpr = style._element.get_or_add_rPr()
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:val"), LANG_META[language]["word_lang"])


def add_footer(section, language: str):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual 03 | Controlled publication QA candidate | ")
    set_run_font(run, language, size=8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def build_docx(source: EditionSource, out_path: Path, image_dir: Path, source_head: str):
    language = source.language
    meta = LANG_META[language]
    doc = Document()
    set_document_defaults(doc, language)
    add_footer(doc.sections[0], language)

    props = doc.core_properties
    props.title = meta["title"]
    props.subject = f"Manual 03 controlled publication-QA artifact; language {language}; source head {source_head}"
    props.author = "Alberto (Al) Leiva"
    props.keywords = "NIST AI RMF, AI 100-1, AI risk management, GOVERN, MAP, MEASURE, MANAGE, GRC"
    props.comments = "Generation and QA do not establish certification, legal compliance, trustworthy-AI achievement, or an audit opinion."

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["title"])
    set_run_font(run, language, size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta["status"])
    set_run_font(run, language, size=10, bold=True)
    control = doc.add_paragraph()
    control.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = control.add_run(
        f"Controlled source revision: {source_head} | Language: {language} | "
        "NIST AI RMF 1.0 / NIST AI 100-1 baseline; version-aware and subject to source-watch controls."
    )
    set_run_font(run, language, size=8.5)

    boundary = doc.add_paragraph()
    boundary.paragraph_format.left_indent = Inches(0.2)
    boundary.paragraph_format.right_indent = Inches(0.2)
    run = boundary.add_run(
        "Assurance boundary: This practical implementation manual does not certify an AI system, "
        "establish legal compliance, prove trustworthy-AI achievement, or provide an audit opinion."
    )
    set_run_font(run, language, size=9, bold=True)

    graphic_counter = [0]
    h = doc.add_paragraph("Implementation paths and operating model", style="Heading 1")
    set_paragraph_keep(h, keep_next=True)
    add_markdown(doc, source.implementation_text, language, image_dir, graphic_counter, "Implementation")
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1], language)
    h = doc.add_paragraph("Controlled 32-chapter manual", style="Heading 1")
    set_paragraph_keep(h, keep_next=True)
    add_markdown(doc, source.chapter_text, language, image_dir, graphic_counter, "Chapter")

    if graphic_counter[0] < 12:
        raise ValueError(f"{language} has too few memory graphics for publication: {graphic_counter[0]}")

    for p in doc.paragraphs:
        for r in p.runs:
            set_run_font(r, language)
    doc.save(out_path)
    return graphic_counter[0]


def convert_pdf(docx_path: Path, pdf_dir: Path) -> Path:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice/soffice not found")
    profile = Path(tempfile.mkdtemp(prefix="manual03-lo-"))
    try:
        command = [
            soffice,
            "--headless",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to", "pdf",
            "--outdir", str(pdf_dir),
            str(docx_path),
        ]
        proc = subprocess.run(command, text=True, capture_output=True, timeout=180)
        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {proc.stdout}\n{proc.stderr}")
        pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
        if not pdf_path.is_file() or pdf_path.stat().st_size < 1000:
            raise RuntimeError(f"PDF was not generated: {pdf_path}")
        return pdf_path
    finally:
        shutil.rmtree(profile, ignore_errors=True)


def inspect_docx(path: Path, language: str, expected_graphics: int) -> dict:
    if not zipfile.is_zipfile(path):
        raise ValueError(f"not a valid DOCX ZIP: {path}")
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    chapter_numbers = [int(x) for x in re.findall(r"(?m)^([0-9]{1,2})\.\s+", text)]
    # Headings may lose their Markdown hash but retain numbering.
    chapter_set = sorted(set(n for n in chapter_numbers if 1 <= n <= 32))
    if len(text) < 30_000:
        raise ValueError(f"DOCX text unexpectedly small: {path}")
    if chapter_set != list(range(1, 33)):
        raise ValueError(f"DOCX chapter inventory incomplete for {language}: {chapter_set}")

    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
        styles = zf.read("word/styles.xml").decode("utf-8", errors="replace")
    alt_count = len(re.findall(r"\bdescr=\"[^\"]+\"", xml))
    if alt_count < expected_graphics:
        raise ValueError(f"DOCX image alt-text count {alt_count} < graphics {expected_graphics}: {path}")
    if LANG_META[language]["word_lang"] not in styles and LANG_META[language]["word_lang"] not in xml:
        raise ValueError(f"DOCX language metadata missing for {language}: {path}")

    return {
        "file": path.name,
        "bytes": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "alt_text_entries": alt_count,
        "chapter_count": len(chapter_set),
        "sha256": sha256(path),
        "status": "PASS",
    }


def inspect_pdf(path: Path, render_dir: Path) -> tuple[dict, list[dict]]:
    pdf = fitz.open(path)
    if pdf.page_count < 10:
        raise ValueError(f"PDF page count unexpectedly small ({pdf.page_count}): {path}")
    render_dir.mkdir(parents=True, exist_ok=True)
    page_rows: list[dict] = []
    blank_pages = []
    for index, page in enumerate(pdf):
        text = page.get_text("text").strip()
        if len(text) < 20:
            blank_pages.append(index + 1)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.35, 1.35), alpha=False)
        png = render_dir / f"page-{index + 1:03d}.png"
        pix.save(png)
        page_rows.append({
            "pdf": path.name,
            "page": index + 1,
            "width_pt": round(page.rect.width, 2),
            "height_pt": round(page.rect.height, 2),
            "text_chars": len(text),
            "render": str(png),
            "automated_status": "PASS" if len(text) >= 20 else "REVIEW",
        })
    if blank_pages:
        raise ValueError(f"possible blank PDF pages in {path.name}: {blank_pages}")
    meta = dict(pdf.metadata or {})
    result = {
        "file": path.name,
        "bytes": path.stat().st_size,
        "pages": pdf.page_count,
        "metadata": meta,
        "sha256": sha256(path),
        "status": "PASS",
    }
    pdf.close()
    return result, page_rows


def make_contact_sheets(render_dirs: dict[str, Path], out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    for language, directory in render_dirs.items():
        pages = sorted(directory.glob("page-*.png"))
        per_sheet = 6
        for sheet_index in range(0, len(pages), per_sheet):
            batch = pages[sheet_index:sheet_index + per_sheet]
            thumbs = []
            for page in batch:
                image = Image.open(page).convert("RGB")
                image.thumbnail((600, 780))
                thumbs.append((page.name, image.copy()))
            cell_w, cell_h = 640, 840
            sheet = Image.new("RGB", (cell_w * 2, cell_h * 3), "white")
            draw = ImageDraw.Draw(sheet)
            label_font = font(18, bold=True)
            for idx, (name, image) in enumerate(thumbs):
                x = (idx % 2) * cell_w + 20
                y = (idx // 2) * cell_h + 35
                draw.text((x, y - 26), name, fill="black", font=label_font)
                sheet.paste(image, (x, y))
            target = out_dir / f"{language.replace('-', '_')}-sheet-{sheet_index // per_sheet + 1:02d}.jpg"
            sheet.save(target, format="JPEG", quality=84, optimize=True)


def write_reports(output_root: Path, sources: dict[str, EditionSource], file_results: dict, page_rows: list[dict], source_head: str):
    qa_dir = output_root / "qa"
    qa_dir.mkdir(parents=True, exist_ok=True)
    manifest = {
        "schema_version": "1.0",
        "manual": "Manual 03 - NIST AI Risk Management Framework Implementation",
        "source_head": source_head,
        "source_branch": "build/nist-ai-rmf-manual-03-2026",
        "controlled_baseline": "NIST AI RMF 1.0 / NIST AI 100-1; NIST AI 600-1 when generative AI is in scope",
        "assurance_boundary": "Publication QA does not establish legal compliance, certification, trustworthy-AI achievement, or an audit opinion.",
        "editions": {},
    }
    for language, source in sources.items():
        manifest["editions"][language] = {
            "chapter_sources": source.chapter_files,
            "implementation_source": source.implementation_file,
            "chapter_count": 32,
            "mermaid_source_blocks": len(MERMAID_RE.findall(source.chapter_text)) + len(MERMAID_RE.findall(source.implementation_text)),
            "artifacts": file_results[language],
        }
    (qa_dir / "MANUAL_03_PUBLICATION_REPORT.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    report_lines = [
        "# Manual 03 Publication Processing Report",
        "",
        f"- Controlled source head: `{source_head}`",
        "- Candidate languages: English (`en`), Latin American Spanish (`es-419`), Brazilian Portuguese (`pt-BR`)",
        "- Automated document-processing status: **PASS**",
        "- Release status: **QA CANDIDATE - human/release controls remain authoritative**",
        "",
        "## Automated checks",
        "",
        "- 32 numbered chapters detected in each DOCX candidate.",
        "- DOCX packages are valid ZIP/OOXML containers.",
        "- Language metadata is present in each DOCX.",
        "- Every generated memory graphic has DOCX alternative text.",
        "- Every PDF page contains extractable text and rendered successfully.",
        "- SHA-256 provenance is recorded for all six publication candidates.",
        "- Rendered page images and contact sheets are included for complete visual review.",
        "",
        "## Assurance boundary",
        "",
        "Successful conversion and automated QA do not constitute human semantic approval, legal compliance, certification, trustworthy-AI achievement, or an audit opinion.",
    ]
    (qa_dir / "MANUAL_03_PUBLICATION_REPORT.md").write_text("\n".join(report_lines) + "\n", encoding="utf-8")

    with (qa_dir / "MANUAL_03_PAGE_QA.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["pdf", "page", "width_pt", "height_pt", "text_chars", "render", "automated_status"])
        writer.writeheader()
        writer.writerows(page_rows)

    candidates = sorted((output_root / "publication").glob("*.docx")) + sorted((output_root / "publication").glob("*.pdf"))
    with (qa_dir / "MANUAL_03_SHA256SUMS.txt").open("w", encoding="utf-8") as handle:
        for path in candidates:
            handle.write(f"{sha256(path)}  {path.name}\n")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-root", type=Path, required=True)
    parser.add_argument("--source-head", default="unknown")
    args = parser.parse_args()

    output_root = args.output_root.resolve()
    publication = output_root / "publication"
    images = output_root / "generated-graphics"
    renders = output_root / "renders"
    contact_sheets = output_root / "contact-sheets"
    for directory in [publication, images, renders, contact_sheets]:
        directory.mkdir(parents=True, exist_ok=True)

    sources = load_sources()
    file_results: dict[str, dict] = {}
    page_rows: list[dict] = []
    render_dirs: dict[str, Path] = {}

    for language, source in sources.items():
        meta = LANG_META[language]
        docx_path = publication / f"{meta['filename']}.docx"
        image_dir = images / language.replace("-", "_")
        image_dir.mkdir(parents=True, exist_ok=True)
        graphic_count = build_docx(source, docx_path, image_dir, args.source_head)
        docx_result = inspect_docx(docx_path, language, graphic_count)
        pdf_path = convert_pdf(docx_path, publication)
        render_dir = renders / language.replace("-", "_")
        pdf_result, rows = inspect_pdf(pdf_path, render_dir)
        page_rows.extend(rows)
        render_dirs[language] = render_dir
        file_results[language] = {
            "graphics": graphic_count,
            "docx": docx_result,
            "pdf": pdf_result,
        }

    make_contact_sheets(render_dirs, contact_sheets)
    write_reports(output_root, sources, file_results, page_rows, args.source_head)
    print(json.dumps({"status": "PASS", "output_root": str(output_root), "languages": list(sources)}, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
