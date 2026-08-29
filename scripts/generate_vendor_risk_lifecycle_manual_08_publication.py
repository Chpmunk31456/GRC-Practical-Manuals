#!/usr/bin/env python3
"""Generate trilingual Manual 08 Vendor and Third-Party Risk publication-QA candidates.

Reuses the established publication-processing core while preserving supplier-risk,
assurance, localization, accessibility, provenance, and human-review boundaries.
Automated generation does not establish vendor acceptability, compliance, assurance,
or eliminate third-party/fourth-party risk.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import generate_ai_security_lifecycle_manual_07_publication as base

ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "07-third-party-risk" / "Vendor_Risk_Lifecycle"
BASELINE = ROOT / ".compliance" / "vendor-risk-lifecycle-manual-08-baseline.json"
core = base.core

LANG_META = {
    "en": {
        "title": "Manual 08 - Vendor and Third-Party Risk Lifecycle",
        "word_lang": "en-US",
        "filename": "Manual_08_Vendor_and_Third_Party_Risk_Lifecycle_EN",
        "status": "CONTROLLED PUBLICATION QA CANDIDATE",
        "implementation_heading": "Implementation paths",
        "chapter_heading": "Controlled 32-chapter manual",
        "control": "Controlled source revision: {source_head} | Language: en | supplier-risk and human decision boundaries retained.",
        "boundary": "Assurance boundary: Risk-based supplier governance reduces uncertainty but does not eliminate third-party or fourth-party risk. Evidence can be incomplete, stale, externally controlled, or subject to material change. Accountable human owners retain approval, exception, residual-risk, renewal, and exit decisions.",
        "figure_title": "Manual 08 memory graphic",
    },
    "es-419": {
        "title": "Manual 08 - Ciclo de Vida de Riesgo de Proveedores y Terceros",
        "word_lang": "es-419",
        "filename": "Manual_08_Vendor_and_Third_Party_Risk_Lifecycle_ES-419",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACION",
        "implementation_heading": "Rutas de implementación",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "control": "Revisión de fuente controlada: {source_head} | Idioma: es-419 | se mantienen los límites de riesgo de proveedores y de decisión humana.",
        "boundary": "Límite de aseguramiento: El gobierno de proveedores basado en riesgo reduce la incertidumbre, pero no elimina el riesgo de terceros o cuartas partes. La evidencia puede ser incompleta, desactualizada, controlada externamente o estar sujeta a cambios materiales. Los responsables humanos conservan las decisiones de aprobación, excepción, riesgo residual, renovación y salida.",
        "figure_title": "Gráfico de memoria del Manual 08",
    },
    "pt-BR": {
        "title": "Manual 08 - Ciclo de Vida de Risco de Fornecedores e Terceiros",
        "word_lang": "pt-BR",
        "filename": "Manual_08_Vendor_and_Third_Party_Risk_Lifecycle_PT-BR",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACAO",
        "implementation_heading": "Caminhos de implementação",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "control": "Revisão da fonte controlada: {source_head} | Idioma: pt-BR | limites de risco de fornecedores e de decisão humana mantidos.",
        "boundary": "Limite de asseguração: A governança de fornecedores baseada em risco reduz a incerteza, mas não elimina o risco de terceiros ou quartas partes. As evidências podem ser incompletas, desatualizadas, controladas externamente ou sujeitas a mudanças materiais. Responsáveis humanos mantêm as decisões de aprovação, exceção, risco residual, renovação e saída.",
        "figure_title": "Gráfico de memória do Manual 08",
    },
}


def source_dir(language: str) -> Path:
    folder = "English" if language == "en" else ("Spanish_es-419" if language == "es-419" else "Portuguese_pt-BR")
    return MANUAL / folder / "source"


def find_localized_chapters(language: str):
    chapters: dict[int, str] = {}
    used: list[str] = []
    for path in sorted(source_dir(language).glob("*.md")):
        found = core.split_chapters(path.read_text(encoding="utf-8"))
        if not found:
            continue
        for number, body in found.items():
            if number in chapters and chapters[number] != body:
                raise ValueError(f"conflicting chapter {number} for {language}: {path}")
            chapters[number] = body
        used.append(str(path.relative_to(ROOT)))
    expected = set(range(1, 33))
    if set(chapters) != expected:
        raise ValueError(f"{language} chapter inventory invalid: {sorted(chapters)}")
    return "\n".join(chapters[n].rstrip() for n in range(1, 33)) + "\n", used


def find_implementation(language: str):
    if language == "en":
        path = MANUAL / "MANUAL_08_IMPLEMENTATION_PATHS.md"
    elif language == "es-419":
        path = source_dir(language) / "RUTAS_DE_IMPLEMENTACION_MANUAL_08.md"
    else:
        path = source_dir(language) / "CAMINHOS_DE_IMPLEMENTACAO_MANUAL_08.md"
    if not path.is_file():
        raise FileNotFoundError(path)
    text = path.read_text(encoding="utf-8")
    if len(core.MERMAID_RE.findall(text)) != 3:
        raise ValueError(f"expected exactly three memory graphics for {language}")
    markers = text.count("**Accessible explanation:**") + text.count("**Explicación accesible:**") + text.count("**Explicação acessível:**")
    if markers != 3:
        raise ValueError(f"expected three accessible graphic explanations for {language}")
    return text, str(path.relative_to(ROOT))


def add_footer(section, language: str):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Manual 08 | {LANG_META[language]['status']} | ")
    core.set_run_font(run, language, size=8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


_base_render = core.render_mermaid_memory_graphic
_base_alt = core.set_image_alt_text


def render_mermaid_memory_graphic(block: str, out_path: Path, title: str) -> str:
    s = str(out_path)
    language = "es-419" if "ES-419" in s or "es_419" in s else ("pt-BR" if "PT-BR" in s or "pt_BR" in s else "en")
    number = title.rsplit(" ", 1)[-1]
    return _base_render(block, out_path, f"{LANG_META[language]['figure_title']} {number}")


def set_image_alt_text(inline_shape, title: str, description: str):
    return _base_alt(inline_shape, title.replace("Manual 03", "Manual 08").replace("Manual 07", "Manual 08"), description.replace("Manual 03", "Manual 08").replace("Manual 07", "Manual 08"))


def build_docx(source: core.EditionSource, out_path: Path, image_dir: Path, source_head: str):
    language = source.language
    meta = LANG_META[language]
    doc = Document()
    core.set_document_defaults(doc, language)
    add_footer(doc.sections[0], language)
    props = doc.core_properties
    props.title = meta["title"]
    props.subject = f"Manual 08 controlled publication-QA artifact; language {language}; source head {source_head}"
    props.author = "Alberto (Al) Leiva"
    props.keywords = "vendor risk, third-party risk, TPRM, C-SCRM, supplier assurance, fourth party, concentration risk, due diligence, monitoring, exit"
    props.comments = "Automated generation does not establish vendor acceptability, compliance, assurance, or elimination of supplier risk."
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["title"])
    core.set_run_font(run, language, size=22, bold=True)
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta["status"]); core.set_run_font(run, language, size=10, bold=True)
    control = doc.add_paragraph(); control.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = control.add_run(meta["control"].format(source_head=source_head))
    core.set_run_font(run, language, size=8.5)
    boundary = doc.add_paragraph(); run = boundary.add_run(meta["boundary"]); core.set_run_font(run, language, size=9, bold=True)
    graphic_counter = [0]
    h = doc.add_paragraph(meta["implementation_heading"], style="Heading 1"); core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.implementation_text, language, image_dir, graphic_counter, "Implementation")
    doc.add_section(WD_SECTION.NEW_PAGE); add_footer(doc.sections[-1], language)
    h = doc.add_paragraph(meta["chapter_heading"], style="Heading 1"); core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.chapter_text, language, image_dir, graphic_counter, "Chapter")
    if graphic_counter[0] != 3:
        raise ValueError(f"{language} expected exactly three memory graphics, got {graphic_counter[0]}")
    for p in doc.paragraphs:
        for r in p.runs:
            core.set_run_font(r, language)
    doc.save(out_path)
    return graphic_counter[0]


def inspect_pdf(path: Path, render_dir: Path):
    pdf = core.fitz.open(path)
    if pdf.page_count < 8:
        raise ValueError(f"PDF page count unexpectedly small ({pdf.page_count}): {path}")
    render_dir.mkdir(parents=True, exist_ok=True)
    rows, blank_pages = [], []
    for index, page in enumerate(pdf):
        text = page.get_text("text").strip()
        if len(text) < 20:
            blank_pages.append(index + 1)
        pix = page.get_pixmap(matrix=core.fitz.Matrix(1.35, 1.35), alpha=False)
        png = render_dir / f"page-{index + 1:03d}.png"; pix.save(png)
        rows.append({"pdf": path.name, "page": index + 1, "width_pt": round(page.rect.width, 2), "height_pt": round(page.rect.height, 2), "text_chars": len(text), "render": str(png), "automated_status": "PASS" if len(text) >= 20 else "REVIEW"})
    if blank_pages:
        raise ValueError(f"possible blank PDF pages in {path.name}: {blank_pages}")
    result = {"file": path.name, "bytes": path.stat().st_size, "pages": pdf.page_count, "metadata": dict(pdf.metadata or {}), "sha256": core.sha256(path), "status": "PASS"}
    pdf.close()
    return result, rows


def write_reports(output_root: Path, sources: dict[str, core.EditionSource], file_results: dict, page_rows: list[dict], source_head: str):
    qa_dir = output_root / "qa"; qa_dir.mkdir(parents=True, exist_ok=True)
    manifest = {"schema_version": "1.0", "manual": "Manual 08 - Vendor and Third-Party Risk Lifecycle", "source_head": source_head, "source_branch": "build/vendor-risk-lifecycle-manual-08-2026", "assurance_boundary": "Automated QA supports document integrity and review; it does not establish vendor acceptability, compliance, assurance, or elimination of third/fourth-party risk.", "editions": {}}
    for language, source in sources.items():
        manifest["editions"][language] = {"chapter_sources": source.chapter_files, "implementation_source": source.implementation_file, "chapter_count": 32, "mermaid_source_blocks": len(core.MERMAID_RE.findall(source.implementation_text)), "artifacts": file_results[language]}
    (qa_dir / "MANUAL_08_PUBLICATION_REPORT.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    (qa_dir / "MANUAL_08_PUBLICATION_REPORT.md").write_text("# Manual 08 Publication Processing Report\n\n" f"- Controlled source head: `{source_head}`\n" "- Candidate languages: English (`en`), Latin American Spanish (`es-419`), Brazilian Portuguese (`pt-BR`)\n" "- Automated document-processing status: **PASS**\n" "- Release status: **QA CANDIDATE - human semantic/accessibility/release controls remain authoritative**\n\n" "## Automated checks\n\n" "- 32 chapters detected in each DOCX candidate.\n" "- Localized implementation-path sources are present for all three editions.\n" "- Supplier-risk and assurance boundaries are retained.\n" "- DOCX packages, language metadata, and image alternative text validated.\n" "- Every PDF page contains extractable text and rendered successfully.\n" "- SHA-256 provenance recorded for all publication candidates.\n" "- Page renders and contact sheets generated for human visual review.\n", encoding="utf-8")
    with (qa_dir / "MANUAL_08_PAGE_QA.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["pdf", "page", "width_pt", "height_pt", "text_chars", "render", "automated_status"]); writer.writeheader(); writer.writerows(page_rows)
    candidates = sorted((output_root / "publication").glob("*.docx")) + sorted((output_root / "publication").glob("*.pdf"))
    with (qa_dir / "MANUAL_08_SHA256SUMS.txt").open("w", encoding="utf-8") as handle:
        for path in candidates: handle.write(f"{core.sha256(path)}  {path.name}\n")

core.MANUAL = MANUAL
core.BASELINE = BASELINE
core.LANG_META = LANG_META
core.find_localized_chapters = find_localized_chapters
core.find_implementation = find_implementation
core.add_footer = add_footer
core.render_mermaid_memory_graphic = render_mermaid_memory_graphic
core.set_image_alt_text = set_image_alt_text
core.build_docx = build_docx
core.inspect_pdf = inspect_pdf
core.write_reports = write_reports

if __name__ == "__main__":
    raise SystemExit(core.main())
