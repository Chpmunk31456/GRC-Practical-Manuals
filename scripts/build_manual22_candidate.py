from pathlib import Path
from docx import Document
from docx.shared import Pt
import re, hashlib, json, subprocess, os

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / '06-cloud-and-technology-risk/Cloud_Security_Controlled_Implementation/controlled'
OUT = ROOT / 'qa/manual22-publication-candidate'
OUT.mkdir(parents=True, exist_ok=True)

sources = {
    'en': BASE / 'en/MANUAL_22_CONTROLLED_EN.md',
    'es-419': BASE / 'es-419/MANUAL_22_CONTROLLED_ES_419.md',
    'pt-BR': BASE / 'pt-BR/MANUAL_22_CONTROLLED_PT_BR.md',
}

names = {
    'en': 'Manual_22_Cloud_Security_Controlled_EN',
    'es-419': 'Manual_22_Cloud_Security_Controlled_ES-419',
    'pt-BR': 'Manual_22_Cloud_Security_Controlled_PT-BR',
}

def add_markdown(doc, text):
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 3))
            continue
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph(); r = p.add_run(line.strip('*')); r.bold = True
            continue
        if line.startswith('`') and line.endswith('`'):
            p = doc.add_paragraph(); p.add_run(line.strip('`')).font.name = 'Liberation Mono'
            continue
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(6)

manifest = {'manual': 22, 'source_commit': os.environ.get('GITHUB_SHA', ''), 'artifacts': []}
for locale, src in sources.items():
    if not src.is_file():
        raise SystemExit(f'missing source: {src}')
    text = src.read_text(encoding='utf-8')
    doc = Document()
    doc.styles['Normal'].font.name = 'Liberation Sans'
    doc.styles['Normal'].font.size = Pt(10.5)
    add_markdown(doc, text)
    stem = names[locale]
    docx = OUT / f'{stem}.docx'
    doc.save(docx)
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', str(OUT), str(docx)], check=True)
    pdf = OUT / f'{stem}.pdf'
    for path in (docx, pdf):
        data = path.read_bytes()
        if not data:
            raise SystemExit(f'empty artifact: {path}')
        manifest['artifacts'].append({'locale': locale, 'file': path.name, 'sha256': hashlib.sha256(data).hexdigest(), 'bytes': len(data)})

(OUT / 'MANUAL_22_CANDIDATE_MANIFEST.json').write_text(json.dumps(manifest, indent=2), encoding='utf-8')
print(json.dumps(manifest, indent=2))
