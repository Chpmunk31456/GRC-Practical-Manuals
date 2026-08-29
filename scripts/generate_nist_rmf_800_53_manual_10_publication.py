#!/usr/bin/env python3
"""Generate trilingual Manual 10 NIST RMF / SP 800-53 publication-QA candidates.

The generator preserves RMF authorization boundaries, SP 800-53 tailoring,
assessment evidence, OSCAL traceability, localization, accessibility,
provenance, and human-review controls. Automated generation does not authorize
a system, certify compliance, establish control effectiveness, or substitute
for an accountable authorizing official or competent assessor.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import generate_nist_csf_2_manual_09_publication as base

ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "01-foundations" / "NIST_RMF_SP_800-53_Controlled_Implementation"
BASELINE = ROOT / ".compliance" / "nist-rmf-800-53-manual-10-baseline.json"
core = base.core

LANG_META = {
    "en": {
        "title": "Manual 10 - NIST RMF and SP 800-53 Controlled Implementation",
        "word_lang": "en-US",
        "filename": "Manual_10_NIST_RMF_SP_800-53_Controlled_Implementation_EN",
        "status": "CONTROLLED PUBLICATION QA CANDIDATE",
        "implementation_heading": "Implementation paths",
        "chapter_heading": "Controlled 32-chapter manual",
        "control": "Controlled source revision: {source_head} | Language: en | RMF authorization and human decision boundaries retained.",
        "boundary": "Assurance boundary: This manual operationalizes NIST RMF and SP 800-53 concepts but does not authorize a system, create NIST certification, prove control effectiveness, establish legal compliance, or replace an accountable authorizing official or competent assessor. Tailoring, assessment, authorization, exceptions, residual risk, and material decisions require accountable human judgment.",
        "figure_title": "Manual 10 memory graphic",
        "figure_caption": "Figure {number}. RMF implementation memory graphic",
    },
    "es-419": {
        "title": "Manual 10 - Implementación Controlada de NIST RMF y SP 800-53",
        "word_lang": "es-419",
        "filename": "Manual_10_NIST_RMF_SP_800-53_Controlled_Implementation_ES-419",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACION",
        "implementation_heading": "Rutas de implementación",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "control": "Revisión de fuente controlada: {source_head} | Idioma: es-419 | se mantienen los límites de autorización RMF y de decisión humana.",
        "boundary": "Límite de aseguramiento: Este manual operacionaliza conceptos de NIST RMF y SP 800-53, pero no autoriza un sistema, no crea certificación NIST, no demuestra efectividad de controles, no establece cumplimiento legal ni sustituye a una autoridad de autorización responsable o a un evaluador competente. La adaptación, evaluación, autorización, excepciones, riesgo residual y decisiones materiales requieren juicio humano responsable.",
        "figure_title": "Gráfico de memoria del Manual 10",
        "figure_caption": "Figura {number}. Gráfico de memoria de implementación RMF",
    },
    "pt-BR": {
        "title": "Manual 10 - Implementação Controlada do NIST RMF e SP 800-53",
        "word_lang": "pt-BR",
        "filename": "Manual_10_NIST_RMF_SP_800-53_Controlled_Implementation_PT-BR",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACAO",
        "implementation_heading": "Caminhos de implementação",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "control": "Revisão da fonte controlada: {source_head} | Idioma: pt-BR | limites de autorização RMF e de decisão humana mantidos.",
        "boundary": "Limite de asseguração: Este manual operacionaliza conceitos do NIST RMF e SP 800-53, mas não autoriza um sistema, não cria certificação NIST, não comprova efetividade de controles, não estabelece conformidade legal nem substitui uma autoridade de autorização responsável ou um avaliador competente. Tailoring, avaliação, autorização, exceções, risco residual e decisões materiais exigem julgamento humano responsável.",
        "figure_title": "Gráfico de memória do Manual 10",
        "figure_caption": "Figura {number}. Gráfico de memória de implementação RMF",
    },
}


def source_dir(language: str) -> Path:
    folder = "English" if language == "en" else ("Spanish_es-419" if language == "es-419" else "Portuguese_pt-BR")
    return MANUAL / folder / "source"


def find_localized_chapters(language: str):
    chapters: dict[int, str] = {}
    used: list[str] = []
    for path in sorted(source_dir(language).glob("*.md")):
        if path.name in {"RUTAS_DE_IMPLEMENTACION_MANUAL_10.md", "CAMINHOS_DE_IMPLEMENTACAO_MANUAL_10.md"}:
            continue
        found = core.split_chapters(path.read_text(encoding="utf-8"))
        if not found:
            continue
        for number, body in found.items():
            if number in chapters and chapters[number] != body:
                raise ValueError(f"conflicting chapter {number} for {language}: {path}")
            chapters[number] = body
        used.append(str(path.relative_to(ROOT)))
    expected = set(range(1, 33))
    if set(chapters) != expected:
        raise ValueError(f"{language} chapter inventory invalid: {sorted(chapters)}")
    return "\n".join(chapters[n].rstrip() for n in range(1, 33)) + "\n", used


def find_implementation(language: str):
    if language == "en":
        path = MANUAL / "MANUAL_10_IMPLEMENTATION_PATHS.md"
    elif language == "es-419":
        path = source_dir(language) / "RUTAS_DE_IMPLEMENTACION_MANUAL_10.md"
    else:
        path = source_dir(language) / "CAMINHOS_DE_IMPLEMENTACAO_MANUAL_10.md"
    if not path.is_file():
        raise FileNotFoundError(path)
    text = path.read_text(encoding="utf-8")
    if len(core.MERMAID_RE.findall(text)) != 3:
        raise ValueError(f"expected exactly three memory graphics for {language}")
    markers = text.count("**Accessible explanation:**") + text.count("**Explicación accesible:**") + text.count("**Explicação acessível:**")
    if markers != 3:
        raise ValueError(f"expected three accessible graphic explanations for {language}")
    for identifier in ("PREPARE", "CATEGORIZE", "SELECT", "IMPLEMENT", "ASSESS", "AUTHORIZE", "MONITOR"):
        if identifier not in text:
            raise ValueError(f"missing RMF step identifier {identifier} in {language} implementation source")
    if "OSCAL" not in text:
        raise ValueError(f"missing OSCAL traceability boundary in {language} implementation source")
    return text, str(path.relative_to(ROOT))


def add_footer(section, language: str):
    existing = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
    if "Manual 10" in existing:
        return
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Manual 10 | {LANG_META[language]['status']} | ")
    core.set_run_font(run, language, size=8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


# This binding is replaced again by the safe runner with the raw Manual 03
# renderer retained by the Manual 06 adapter. Keeping it module-local prevents
# Manual 09 labels from becoming part of Manual 10 artifacts.
_base_render = base._base_render
_base_alt = base._base_alt


def render_mermaid_memory_graphic(block: str, out_path: Path, title: str) -> str:
    s = str(out_path)
    language = "es-419" if "ES-419" in s or "es_419" in s else ("pt-BR" if "PT-BR" in s or "pt_BR" in s else "en")
    number = title.rsplit(" ", 1)[-1]
    return _base_render(block, out_path, f"{LANG_META[language]['figure_title']} {number}")


def set_image_alt_text(inline_shape, title: str, description: str):
    normalized_title = title
    normalized_description = description
    for prior in ("Manual 03", "Manual 06", "Manual 07", "Manual 09"):
        normalized_title = normalized_title.replace(prior, "Manual 10")
        normalized_description = normalized_description.replace(prior, "Manual 10")
    return _base_alt(inline_shape, normalized_title, normalized_description)


def localize_figure_captions(doc: Document, language: str) -> None:
    template = LANG_META[language]["figure_caption"]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text.startswith("Figure ") or "Implementation memory graphic" not in text:
            continue
        number = text.split(".", 1)[0].split()[-1]
        paragraph.text = template.format(number=number)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        core.set_paragraph_keep(paragraph, keep_next=True)


def build_docx(source: core.EditionSource, out_path: Path, image_dir: Path, source_head: str):
    language = source.language
    meta = LANG_META[language]
    doc = Document()
    core.set_document_defaults(doc, language)
    add_footer(doc.sections[0], language)
    props = doc.core_properties
    props.title = meta["title"]
    props.subject = f"Manual 10 controlled publication-QA artifact; language {language}; source head {source_head}"
    props.author = "Alberto (Al) Leiva"
    props.keywords = "NIST RMF, SP 800-53, SP 800-53A, SP 800-53B, SP 800-18, OSCAL, authorization, continuous monitoring, evidence"
    props.comments = "Automated generation does not authorize a system, establish control effectiveness, create certification, or replace accountable human authorization and assessment decisions."

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["title"])
    core.set_run_font(run, language, size=22, bold=True)
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta["status"]); core.set_run_font(run, language, size=10, bold=True)
    control = doc.add_paragraph(); control.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = control.add_run(meta["control"].format(source_head=source_head))
    core.set_run_font(run, language, size=8.5)
    boundary = doc.add_paragraph(); run = boundary.add_run(meta["boundary"]); core.set_run_font(run, language, size=9, bold=True)

    graphic_counter = [0]
    h = doc.add_paragraph(meta["implementation_heading"], style="Heading 1"); core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.implementation_text, language, image_dir, graphic_counter, "Implementation")
    doc.add_section(WD_SECTION.NEW_PAGE); add_footer(doc.sections[-1], language)
    h = doc.add_paragraph(meta["chapter_heading"], style="Heading 1"); core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.chapter_text, language, image_dir, graphic_counter, "Chapter")
    if graphic_counter[0] != 3:
        raise ValueError(f"{language} expected exactly three memory graphics, got {graphic_counter[0]}")
    localize_figure_captions(doc, language)
    for p in doc.paragraphs:
        for r in p.runs:
            core.set_run_font(r, language)
    doc.save(out_path)
    return graphic_counter[0]


def write_reports(output_root: Path, sources: dict[str, core.EditionSource], file_results: dict, page_rows: list[dict], source_head: str):
    qa_dir = output_root / "qa"; qa_dir.mkdir(parents=True, exist_ok=True)
    manifest = {
        "schema_version": "1.0",
        "manual": "Manual 10 - NIST RMF and SP 800-53 Controlled Implementation",
        "source_head": source_head,
        "source_branch": "build/nist-rmf-800-53-manual-10-2026",
        "assurance_boundary": "Automated QA supports document integrity and review; it does not authorize a system, establish control effectiveness, create certification, or replace accountable human assessment and authorization decisions.",
        "editions": {},
    }
    for language, source in sources.items():
        manifest["editions"][language] = {
            "chapter_sources": source.chapter_files,
            "implementation_source": source.implementation_file,
            "chapter_count": 32,
            "mermaid_source_blocks": len(core.MERMAID_RE.findall(source.implementation_text)),
            "artifacts": file_results[language],
        }
    (qa_dir / "MANUAL_10_PUBLICATION_REPORT.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    (qa_dir / "MANUAL_10_PUBLICATION_REPORT.md").write_text(
        "# Manual 10 Publication Processing Report\n\n"
        f"- Controlled source head: `{source_head}`\n"
        "- Candidate languages: English (`en`), Latin American Spanish (`es-419`), Brazilian Portuguese (`pt-BR`)\n"
        "- Automated document-processing status: **PASS**\n"
        "- Release status: **QA CANDIDATE - human semantic/accessibility/release controls remain authoritative**\n\n"
        "## Automated checks\n\n"
        "- 32 chapters detected in each DOCX candidate.\n"
        "- Three localized RMF implementation graphics and accessible explanations are present in each edition.\n"
        "- PREPARE, CATEGORIZE, SELECT, IMPLEMENT, ASSESS, AUTHORIZE, and MONITOR identifiers are retained.\n"
        "- OSCAL traceability boundary is retained.\n"
        "- DOCX packages, language metadata, and image alternative text validated.\n"
        "- Every PDF page contains extractable text and rendered successfully.\n"
        "- SHA-256 provenance recorded for all publication candidates.\n"
        "- Page renders and contact sheets generated for human visual review.\n\n"
        "Automated QA does not authorize a system, prove control effectiveness, establish certification, or replace accountable human assessment and authorization decisions.\n",
        encoding="utf-8",
    )
    with (qa_dir / "MANUAL_10_PAGE_QA.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["pdf", "page", "width_pt", "height_pt", "text_chars", "render", "automated_status"])
        writer.writeheader(); writer.writerows(page_rows)
    candidates = sorted((output_root / "publication").glob("*.docx")) + sorted((output_root / "publication").glob("*.pdf"))
    with (qa_dir / "MANUAL_10_SHA256SUMS.txt").open("w", encoding="utf-8") as handle:
        for path in candidates:
            handle.write(f"{core.sha256(path)}  {path.name}\n")


core.MANUAL = MANUAL
core.BASELINE = BASELINE
core.LANG_META = LANG_META
core.find_localized_chapters = find_localized_chapters
core.find_implementation = find_implementation
core.add_footer = add_footer
core.render_mermaid_memory_graphic = render_mermaid_memory_graphic
core.set_image_alt_text = set_image_alt_text
core.build_docx = build_docx
core.write_reports = write_reports

if __name__ == "__main__":
    raise SystemExit(core.main())
