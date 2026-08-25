#!/usr/bin/env python3
"""Replace one-cell layout tables with accessible callout paragraphs.

This targeted remediation is intentionally limited to the controlled English
ISO/IEC 42001 source. Data tables remain unchanged, including their existing
header-row metadata and geometry.
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


EXPECTED_CALLOUT_TABLES = 8


def set_attribute(element: OxmlElement, name: str, value: str) -> None:
    element.set(qn(name), value)


def add_callout_format(paragraph_xml: OxmlElement, index: int, total: int) -> None:
    paragraph_properties = paragraph_xml.get_or_add_pPr()

    shading = paragraph_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        paragraph_properties.append(shading)
    set_attribute(shading, "w:val", "clear")
    set_attribute(shading, "w:color", "auto")
    set_attribute(shading, "w:fill", "FFF4DA")

    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is not None:
        paragraph_properties.remove(borders)
    borders = OxmlElement("w:pBdr")
    sides = ["left", "right"]
    if index == 0:
        sides.append("top")
    if index == total - 1:
        sides.append("bottom")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        set_attribute(border, "w:val", "single")
        set_attribute(border, "w:sz", "8")
        set_attribute(border, "w:space", "4")
        set_attribute(border, "w:color", "D5B46B")
        borders.append(border)
    paragraph_properties.append(borders)

    indentation = paragraph_properties.find(qn("w:ind"))
    if indentation is None:
        indentation = OxmlElement("w:ind")
        paragraph_properties.append(indentation)
    set_attribute(indentation, "w:left", "120")
    set_attribute(indentation, "w:right", "120")

    spacing = paragraph_properties.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        paragraph_properties.append(spacing)
    set_attribute(spacing, "w:before", "80" if index == 0 else "0")
    set_attribute(spacing, "w:after", "80" if index == total - 1 else "0")


def remediate(input_path: Path, output_path: Path) -> None:
    document = Document(input_path)
    callouts = [table for table in document.tables if len(table.rows) == 1 and len(table.columns) == 1]
    if len(callouts) != EXPECTED_CALLOUT_TABLES:
        raise ValueError(
            f"expected {EXPECTED_CALLOUT_TABLES} one-cell callout tables, found {len(callouts)}"
        )

    for table in callouts:
        table_xml = table._tbl
        cell_paragraphs = list(table.cell(0, 0)._tc.p_lst)
        if not cell_paragraphs:
            raise ValueError("callout table contains no paragraphs")
        for index, paragraph in enumerate(cell_paragraphs):
            copied = deepcopy(paragraph)
            add_callout_format(copied, index, len(cell_paragraphs))
            table_xml.addprevious(copied)
        table_xml.getparent().remove(table_xml)

    remaining_callouts = [
        table for table in document.tables if len(table.rows) == 1 and len(table.columns) == 1
    ]
    if remaining_callouts:
        raise ValueError("one or more layout tables remain after remediation")
    if len(document.tables) != 33:
        raise ValueError(f"expected 33 data tables after remediation, found {len(document.tables)}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    remediate(args.input, args.output)


if __name__ == "__main__":
    main()
