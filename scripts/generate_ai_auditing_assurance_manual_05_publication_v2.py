#!/usr/bin/env python3
"""Manual 05 publication generator v2 with source-relative DOCX completeness QA.

This wrapper preserves all v1 controls and replaces only the arbitrary fixed
DOCX character floor with a fail-closed comparison against the visible source
content for the same language.
"""
from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document

import generate_ai_auditing_assurance_manual_05_publication as v1


def _visible_source_text(language: str) -> str:
    chapters, _ = v1.find_localized_chapters(language)
    implementation, _ = v1.find_implementation(language)
    source = implementation + "\n" + chapters
    # Mermaid source is rendered as a figure and should not inflate the expected
    # body-text count. Markdown punctuation is likewise not expected in DOCX text.
    source = re.sub(r"(?ms)^```mermaid\s*\n.*?^```\s*$", "", source)
    source = re.sub(r"(?m)^#{1,6}\s*", "", source)
    source = re.sub(r"[*_`]+", "", source)
    source = re.sub(r"\[(.*?)\]\([^)]*\)", r"\1", source)
    source = re.sub(r"\s+", " ", source).strip()
    return source


def inspect_docx(path: Path, language: str, expected_graphics: int) -> dict:
    if not zipfile.is_zipfile(path):
        raise ValueError(f"not a valid DOCX ZIP: {path}")
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    normalized_text = re.sub(r"\s+", " ", text).strip()

    chapter_numbers = [
        int(x)
        for x in re.findall(
            r"(?mi)^(?:Chapter|Cap[ií]tulo)\s+([0-9]{1,2})\s+[—-]", text
        )
    ]
    chapter_set = sorted(set(n for n in chapter_numbers if 1 <= n <= 32))
    if chapter_set != list(range(1, 33)):
        raise ValueError(
            f"DOCX chapter inventory incomplete for {language}: {chapter_set}"
        )

    visible_source = _visible_source_text(language)
    # A rendered DOCX omits Markdown syntax and Mermaid code but should retain
    # substantially all narrative content. Require at least 80% of normalized
    # visible-source characters, with a hard minimum of 8,000 characters.
    required_chars = max(8000, int(len(visible_source) * 0.80))
    if len(normalized_text) < required_chars:
        raise ValueError(
            f"DOCX narrative completeness below threshold for {language}: "
            f"{len(normalized_text)} < {required_chars} normalized characters"
        )

    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
        styles = zf.read("word/styles.xml").decode("utf-8", errors="replace")
    alt_count = len(re.findall(r"\bdescr=\"[^\"]+\"", xml))
    if alt_count < expected_graphics:
        raise ValueError(
            f"DOCX image alt-text count {alt_count} < graphics "
            f"{expected_graphics}: {path}"
        )
    if (
        v1.LANG_META[language]["word_lang"] not in styles
        and v1.LANG_META[language]["word_lang"] not in xml
    ):
        raise ValueError(f"DOCX language metadata missing for {language}: {path}")

    return {
        "file": path.name,
        "bytes": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "alt_text_entries": alt_count,
        "chapter_count": len(chapter_set),
        "normalized_text_chars": len(normalized_text),
        "source_completeness_floor_chars": required_chars,
        "sha256": v1.core.sha256(path),
        "status": "PASS",
    }


v1.inspect_docx = inspect_docx
v1.core.inspect_docx = inspect_docx

if __name__ == "__main__":
    raise SystemExit(v1.core.main())
