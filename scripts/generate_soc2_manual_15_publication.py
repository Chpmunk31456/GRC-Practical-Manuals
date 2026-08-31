#!/usr/bin/env python3
from __future__ import annotations
import hashlib, json, re, subprocess
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
MANUAL=ROOT/'04-regulatory-compliance'/'SOC_2_Controlled_Implementation'
PUB=MANUAL/'publication'
QA=MANUAL/'qa'
LANGS={
'en':('English/source/CONTROLLED_CHAPTERS_01_32.md','Manual 15 - SOC 2 Controlled Implementation','Manual_15_SOC_2_Controlled_Implementation_EN','en-US'),
'es-419':('Spanish_es-419/source/CAPITULOS_CONTROLADOS_01_32.md','Manual 15 - Implementación Controlada de SOC 2','Manual_15_SOC_2_Controlled_Implementation_ES-419','es-419'),
'pt-BR':('Portuguese_pt-BR/source/CAPITULOS_CONTROLADOS_01_32.md','Manual 15 - Implementação Controlada de SOC 2','Manual_15_SOC_2_Controlled_Implementation_PT-BR','pt-BR')}
CHAPTER_RE=re.compile(r'^## Chapter\s+(\d{2})\s+—\s+(.+)$',re.M)

def sha256(p):
 h=hashlib.sha256()
 with open(p,'rb') as f:
  for c in iter(lambda:f.read(1024*1024),b''): h.update(c)
 return h.hexdigest()

def add_page_number(section):
 p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 p.add_run('Manual 15 | CONTROLLED PUBLICATION CANDIDATE | ')
 fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); p._p.append(fld)

def add_markdown(doc,text):
 for raw in text.splitlines():
  line=raw.rstrip()
  if line.startswith('# '): doc.add_heading(line[2:].strip(),0)
  elif line.startswith('## '): doc.add_heading(line[3:].strip(),1)
  elif line.startswith('### '): doc.add_heading(line[4:].strip(),2)
  elif line.startswith('- '): doc.add_paragraph(line[2:].strip(),style='List Bullet')
  elif line.strip(): doc.add_paragraph(line.strip())

def build(lang,rel,title,stem,word_lang):
 src=MANUAL/rel; text=src.read_text(encoding='utf-8')
 nums=[int(x[0]) for x in CHAPTER_RE.findall(text)]
 if nums!=list(range(1,33)): raise SystemExit(f'{lang}: chapter inventory invalid: {nums}')
 out=PUB/lang; out.mkdir(parents=True,exist_ok=True)
 doc=Document(); add_page_number(doc.sections[0])
 props=doc.core_properties; props.title=title; props.subject='Controlled SOC 2 implementation and readiness guidance'; props.author='GRC Practical Manuals Project'
 t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; t.add_run(title).bold=True
 b=doc.add_paragraph(); b.add_run('Authority boundary: original implementation guidance. SOC 2 is an independent CPA attestation examination, not a certification. Localized editions are controlled implementation translations and are not AICPA-authorized translations. The controlled English edition governs interpretation.').bold=True
 handoff=(MANUAL/'MANUAL_15_LOCALIZATION_AND_ARTIFACT_HANDOFF.md').read_text(encoding='utf-8')
 add_markdown(doc,handoff); doc.add_page_break(); add_markdown(doc,text)
 docx=out/f'{stem}.docx'; doc.save(docx)
 subprocess.run(['libreoffice','--headless','--convert-to','pdf','--outdir',str(out),str(docx)],check=True,stdout=subprocess.PIPE,stderr=subprocess.PIPE)
 pdf=out/f'{stem}.pdf'
 if not pdf.is_file(): raise SystemExit(f'missing PDF {pdf}')
 return {'language':lang,'source':str(src.relative_to(ROOT)),'source_sha256':sha256(src),'docx':str(docx.relative_to(ROOT)),'pdf':str(pdf.relative_to(ROOT)),'docx_sha256':sha256(docx),'pdf_sha256':sha256(pdf)}

def main():
 PUB.mkdir(parents=True,exist_ok=True); QA.mkdir(parents=True,exist_ok=True)
 results=[build(lang,*meta) for lang,meta in LANGS.items()]
 report={'schema_version':'1.0','manual':'Manual 15 - SOC 2 Controlled Implementation','release_state':'publication-candidate','languages':['en','es-419','pt-BR'],'editions':results}
 (QA/'MANUAL_15_PUBLICATION_REPORT.json').write_text(json.dumps(report,indent=2,ensure_ascii=False)+'\n',encoding='utf-8')
 with (QA/'MANUAL_15_SHA256SUMS.txt').open('w',encoding='utf-8') as h:
  for r in results:
   h.write(f"{r['docx_sha256']}  {Path(r['docx']).name}\n{r['pdf_sha256']}  {Path(r['pdf']).name}\n")
 print(json.dumps(report,indent=2,ensure_ascii=False))
if __name__=='__main__': main()
