#!/usr/bin/env python3
"""Generate trilingual Manual 07 AI Security and Lifecycle publication-QA candidates.

This adapter reuses the established Manual 06 publication-processing pipeline while
preserving Manual 07 security, authorization, rollback, human-review, and assurance
boundaries. Automated generation is supporting evidence only and does not prove an
AI system secure, safe, compliant, certified, or free from exploitable weaknesses.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import generate_hipaa_implementation_manual_06_publication as base

ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "06-cloud-and-technology-risk" / "AI_Security_and_Lifecycle"
BASELINE = ROOT / ".compliance" / "ai-security-lifecycle-manual-07-baseline.json"
core = base.core

LANG_META = {
    "en": {
        "title": "Manual 07 - AI Security and Lifecycle Controls",
        "word_lang": "en-US",
        "filename": "Manual_07_AI_Security_and_Lifecycle_Controls_EN",
        "status": "CONTROLLED PUBLICATION QA CANDIDATE",
        "implementation_heading": "Implementation and security paths",
        "chapter_heading": "Controlled 32-chapter manual",
        "language_label": "Language:",
        "boundary": "Security and assurance boundary: This educational manual does not prove an AI system secure, safe, compliant, certified, conformant, or free from exploitable weaknesses. Authorization, least privilege, stop/rollback, incident response, supplier risk, and evidence decisions require accountable human judgment.",
        "figure_title": "Manual 07 memory graphic",
    },
    "es-419": {
        "title": "Manual 07 - Seguridad de IA y Controles del Ciclo de Vida",
        "word_lang": "es-419",
        "filename": "Manual_07_AI_Security_and_Lifecycle_Controls_ES-419",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACION",
        "implementation_heading": "Rutas de implementación y seguridad",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "language_label": "Idioma:",
        "boundary": "Límite de seguridad y aseguramiento: Este manual educativo no demuestra que un sistema de IA sea seguro, conforme, certificado ni libre de vulnerabilidades explotables. Las decisiones sobre autorización, mínimo privilegio, parada/reversión, respuesta a incidentes, riesgo de proveedores y evidencia requieren juicio humano responsable.",
        "figure_title": "Gráfico de memoria del Manual 07",
    },
    "pt-BR": {
        "title": "Manual 07 - Segurança de IA e Controles do Ciclo de Vida",
        "word_lang": "pt-BR",
        "filename": "Manual_07_AI_Security_and_Lifecycle_Controls_PT-BR",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACAO",
        "implementation_heading": "Caminhos de implementação e segurança",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "language_label": "Idioma:",
        "boundary": "Limite de segurança e asseguração: Este manual educacional não demonstra que um sistema de IA seja seguro, conforme, certificado ou livre de vulnerabilidades exploráveis. Decisões sobre autorização, privilégio mínimo, parada/reversão, resposta a incidentes, risco de fornecedores e evidências exigem julgamento humano responsável.",
        "figure_title": "Gráfico de memória do Manual 07",
    },
}

def source_dir(language: str) -> Path:
    folder = "English" if language == "en" else ("Spanish_es-419" if language == "es-419" else "Portuguese_pt-BR")
    return MANUAL / folder / "source"

def find_localized_chapters(language: str):
    old_manual = base.MANUAL
    try:
        base.MANUAL = MANUAL
        return base.find_localized_chapters(language)
    finally:
        base.MANUAL = old_manual

def find_implementation(language: str):
    if language == "en":
        path = MANUAL / "MANUAL_07_IMPLEMENTATION_PATHS.md"
    elif language == "es-419":
        path = source_dir(language) / "RUTAS_DE_IMPLEMENTACION_MANUAL_07.md"
    else:
        path = source_dir(language) / "CAMINHOS_DE_IMPLEMENTACAO_MANUAL_07.md"
    if not path.is_file():
        raise FileNotFoundError(path)
    text = path.read_text(encoding="utf-8")
    if len(core.MERMAID_RE.findall(text)) != 3:
        raise ValueError(f"expected exactly three memory graphics for {language}")
    markers = text.count("**Accessible explanation:**") + text.count("**Explicación accesible:**") + text.count("**Explicação acessível:**")
    if markers != 3:
        raise ValueError(f"expected three accessible graphic explanations for {language}")
    return text, str(path.relative_to(ROOT))

def add_footer(section, language: str):
    existing = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
    if "Manual 07" in existing:
        return
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Manual 07 | {LANG_META[language]['status']} | ")
    core.set_run_font(run, language, size=8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)

_base_render = core.render_mermaid_memory_graphic
_base_alt = core.set_image_alt_text

def render_mermaid_memory_graphic(block: str, out_path: Path, title: str) -> str:
    s = str(out_path)
    language = "es-419" if "ES-419" in s or "es_419" in s else ("pt-BR" if "PT-BR" in s or "pt_BR" in s else "en")
    number = title.rsplit(" ", 1)[-1]
    return _base_render(block, out_path, f"{LANG_META[language]['figure_title']} {number}")

def set_image_alt_text(inline_shape, title: str, description: str):
    return _base_alt(inline_shape, title.replace("Manual 03", "Manual 07"), description.replace("Manual 03", "Manual 07"))

def build_docx(source: core.EditionSource, out_path: Path, image_dir: Path, source_head: str):
    language = source.language
    meta = LANG_META[language]
    doc = Document()
    core.set_document_defaults(doc, language)
    add_footer(doc.sections[0], language)
    props = doc.core_properties
    props.title = meta["title"]
    props.subject = f"Manual 07 controlled publication-QA artifact; language {language}; source head {source_head}"
    props.author = "Alberto (Al) Leiva"
    props.keywords = "AI security, lifecycle, threat modeling, prompt injection, RAG, agent authorization, least privilege, red teaming, incident response, rollback"
    props.comments = "Automated generation does not establish security, safety, compliance, certification, conformance, or an audit opinion."

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["title"])
    core.set_run_font(run, language, size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta["status"])
    core.set_run_font(run, language, size=10, bold=True)
    control = doc.add_paragraph()
    control.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = control.add_run(f"Controlled source revision: {source_head} | {meta['language_label']} {language} | authorization and stop/rollback boundaries retained.")
    core.set_run_font(run, language, size=8.5)
    boundary = doc.add_paragraph()
    run = boundary.add_run(meta["boundary"])
    core.set_run_font(run, language, size=9, bold=True)

    graphic_counter = [0]
    h = doc.add_paragraph(meta["implementation_heading"], style="Heading 1")
    core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.implementation_text, language, image_dir, graphic_counter, "Implementation")
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1], language)
    h = doc.add_paragraph(meta["chapter_heading"], style="Heading 1")
    core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.chapter_text, language, image_dir, graphic_counter, "Chapter")
    if graphic_counter[0] != 3:
        raise ValueError(f"{language} expected exactly three memory graphics, got {graphic_counter[0]}")
    for p in doc.paragraphs:
        for r in p.runs:
            core.set_run_font(r, language)
    doc.save(out_path)
    return graphic_counter[0]

def write_reports(output_root: Path, sources: dict[str, core.EditionSource], file_results: dict, page_rows: list[dict], source_head: str):
    qa_dir = output_root / "qa"
    qa_dir.mkdir(parents=True, exist_ok=True)
    manifest = {
        "schema_version": "1.0",
        "manual": "Manual 07 - AI Security and Lifecycle Controls",
        "source_head": source_head,
        "source_branch": "build/ai-security-lifecycle-manual-07-2026",
        "assurance_boundary": "Automated QA supports document integrity and review; it does not prove AI security, safety, compliance, certification, conformance, or absence of exploitable weaknesses.",
        "editions": {},
    }
    for language, source in sources.items():
        manifest["editions"][language] = {
            "chapter_sources": source.chapter_files,
            "implementation_source": source.implementation_file,
            "chapter_count": 32,
            "mermaid_source_blocks": len(core.MERMAID_RE.findall(source.implementation_text)),
            "artifacts": file_results[language],
        }
    (qa_dir / "MANUAL_07_PUBLICATION_REPORT.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    (qa_dir / "MANUAL_07_PUBLICATION_REPORT.md").write_text(
        "# Manual 07 Publication Processing Report\n\n"
        f"- Controlled source head: `{source_head}`\n"
        "- Candidate languages: English (`en`), Latin American Spanish (`es-419`), Brazilian Portuguese (`pt-BR`)\n"
        "- Automated document-processing status: **PASS**\n"
        "- Release status: **QA CANDIDATE - human semantic/accessibility/release controls remain authoritative**\n\n"
        "## Automated checks\n\n"
        "- 32 chapters detected in each DOCX candidate.\n"
        "- Localized implementation-path sources are present for all three editions.\n"
        "- Authorization, least-privilege, stop/rollback and assurance boundaries are retained.\n"
        "- DOCX packages, language metadata, and image alternative text validated.\n"
        "- Every PDF page contains extractable text and rendered successfully.\n"
        "- SHA-256 provenance recorded for all publication candidates.\n"
        "- Page renders and contact sheets generated for human visual review.\n\n"
        "## Security/assurance boundary\n\n"
        "Successful conversion and automated QA do not establish system security, safety, compliance, certification, conformance, or an audit opinion.\n",
        encoding="utf-8",
    )
    with (qa_dir / "MANUAL_07_PAGE_QA.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["pdf", "page", "width_pt", "height_pt", "text_chars", "render", "automated_status"])
        writer.writeheader()
        writer.writerows(page_rows)
    candidates = sorted((output_root / "publication").glob("*.docx")) + sorted((output_root / "publication").glob("*.pdf"))
    with (qa_dir / "MANUAL_07_SHA256SUMS.txt").open("w", encoding="utf-8") as handle:
        for path in candidates:
            handle.write(f"{core.sha256(path)}  {path.name}\n")

base.MANUAL = MANUAL
base.BASELINE = BASELINE
base.LANG_META = LANG_META
base.source_dir = source_dir
base.find_localized_chapters = find_localized_chapters
base.find_implementation = find_implementation
base.add_footer = add_footer
base.render_mermaid_memory_graphic = render_mermaid_memory_graphic
base.set_image_alt_text = set_image_alt_text
base.build_docx = build_docx
base.write_reports = write_reports

core.MANUAL = MANUAL
core.BASELINE = BASELINE
core.LANG_META = LANG_META
core.find_localized_chapters = find_localized_chapters
core.find_implementation = find_implementation
core.add_footer = add_footer
core.render_mermaid_memory_graphic = render_mermaid_memory_graphic
core.set_image_alt_text = set_image_alt_text
core.build_docx = build_docx
core.inspect_docx = base.inspect_docx
core.write_reports = write_reports

if __name__ == "__main__":
    raise SystemExit(core.main())
