from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import sys
import zipfile
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
TOOLS = {
    "GRC_Program_Charter.csv": 22,
    "GRC_Operating_Model_and_RACI.csv": 18,
    "Annual_GRC_Plan_and_Calendar.csv": 19,
    "GRC_Initiative_Portfolio_Tracker.csv": 24,
    "Governance_Meeting_and_Decision_Log.csv": 20,
    "GRC_Maturity_Assessment.csv": 20,
}

EDITIONS = [
    {
        "locale": "en",
        "md": ROOT / "English/GRC_Program_Management_and_Operating_Model_Toolkit_English_v1.0.md",
        "qa": ROOT / "English",
        "title": "GRC Program Management and Operating Model Toolkit",
        "status": "Controlled English master",
    },
    {
        "locale": "es-419",
        "md": ROOT / "translations/es-419/Kit_de_Gestion_del_Programa_GRC_y_Modelo_Operativo_es-419_v1.0.md",
        "qa": ROOT / "translations/es-419",
        "title": "Kit de Gestión del Programa GRC y Modelo Operativo",
        "status": "Candidato de publicación con traducción asistida por máquina",
    },
    {
        "locale": "pt-BR",
        "md": ROOT / "translations/pt-BR/Kit_de_Gestao_do_Programa_GRC_e_Modelo_Operacional_pt-BR_v1.0.md",
        "qa": ROOT / "translations/pt-BR",
        "title": "Kit de Gestão do Programa GRC e Modelo Operacional",
        "status": "Candidato a publicação com tradução assistida por máquina",
    },
]


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    if "Document Metadata" not in styles:
        meta = styles.add_style("Document Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["Document Metadata"]
    meta.font.name = "Calibri"
    meta.font.size = Pt(10)
    meta.font.color.rgb = RGBColor(89, 89, 89)
    meta.paragraph_format.space_after = Pt(4)


def add_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(31, 77, 120)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def markdown_body(text: str) -> list[str]:
    if text.startswith("---\n"):
        _, _, rest = text.partition("---\n")
        _, sep, body = rest.partition("\n---\n")
        if not sep:
            raise ValueError("YAML boundary missing")
        return body.lstrip("\n").splitlines()
    return text.splitlines()


def build_docx(md_path: Path, title: str, status: str) -> Path:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "GRC Program Management and Operating Model | Repository 1.11"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    add_page_field(section.footer.paragraphs[0])

    first_h1 = True
    for raw in markdown_body(text):
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            heading = line[2:].strip()
            if first_h1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(80)
                p.paragraph_format.space_after = Pt(12)
                r = p.add_run(heading)
                r.font.name = "Calibri"
                r.font.size = Pt(28)
                r.font.bold = True
                r.font.color.rgb = RGBColor(11, 37, 69)
                p2 = doc.add_paragraph(style="Document Metadata")
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.add_run(status)
                p3 = doc.add_paragraph(style="Document Metadata")
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p3.add_run("Repository release 1.11 | 1 August 2026")
                first_h1 = False
            else:
                doc.add_paragraph(heading, style="Heading 1")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:].strip())
        elif line.startswith("> "):
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(6.5)
            cell = table.cell(0, 0)
            cell.width = Inches(6.5)
            set_cell_margins(cell)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F4F6F9")
            cell._tc.get_or_add_tcPr().append(shading)
            add_inline(cell.paragraphs[0], line[2:].strip())
        else:
            p = doc.add_paragraph()
            add_inline(p, line)

    props = doc.core_properties
    props.title = title
    props.author = "Alberto Al Leiva"
    props.subject = "GRC program management and operating model"
    props.keywords = "GRC, governance, risk, compliance, operating model"
    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


def pdf_inline(text: str) -> str:
    text = escape(text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`([^`]+)`", r'<font name="Courier" color="#1F4D78">\1</font>', text)
    return text


def convert_pdf(docx: Path) -> Path:
    md = docx.with_suffix(".md")
    pdf = docx.with_suffix(".pdf")
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10,
                          leading=13, spaceAfter=6, textColor=colors.HexColor("#202124"))
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16,
                        leading=19, spaceBefore=14, spaceAfter=8, textColor=colors.HexColor("#2E74B5"), keepWithNext=True)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12,
                        leading=15, spaceBefore=10, spaceAfter=6, textColor=colors.HexColor("#2E74B5"), keepWithNext=True)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=22, firstLineIndent=-10, bulletIndent=10, spaceAfter=4)
    title = ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24,
                           leading=28, alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"), spaceAfter=16)
    story = []
    first_h1 = True
    for raw in markdown_body(md.read_text(encoding="utf-8")):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            if first_h1:
                story.extend([Spacer(1, 1.1 * inch), Paragraph(pdf_inline(line[2:]), title),
                              Paragraph("Repository release 1.11 | 1 August 2026", ParagraphStyle("Meta", parent=body, alignment=TA_CENTER, textColor=colors.grey)),
                              Spacer(1, 0.65 * inch)])
                first_h1 = False
            else:
                story.append(PageBreak())
                story.append(Paragraph(pdf_inline(line[2:]), h1))
        elif line.startswith("## "):
            story.append(Paragraph(pdf_inline(line[3:]), h2))
        elif line.startswith("- "):
            story.append(Paragraph(pdf_inline(line[2:]), bullet, bulletText="•"))
        else:
            story.append(Paragraph(pdf_inline(line.lstrip("> ")), body))

    def page(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(inch, 0.55 * inch, "GRC Program Management and Operating Model | Repository 1.11")
        canvas.drawRightString(7.5 * inch, 0.55 * inch, f"Page {document.page}")
        canvas.restoreState()

    document = SimpleDocTemplate(str(pdf), pagesize=letter, leftMargin=inch, rightMargin=inch,
                                 topMargin=0.8 * inch, bottomMargin=0.8 * inch,
                                 title=md.stem, author="Alberto Al Leiva")
    document.build(story, onFirstPage=page, onLaterPages=page)
    if not pdf.is_file() or pdf.stat().st_size == 0:
        raise RuntimeError(f"PDF generation failed: {pdf}")
    return pdf


def validate_tools() -> dict[str, int]:
    counts = {}
    for name, expected in TOOLS.items():
        path = ROOT / "tools" / name
        with path.open(encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.reader(handle))
        if len(rows) != 1:
            raise ValueError(f"{name}: expected header-only template, got {len(rows)} rows")
        actual = len(rows[0])
        if actual != expected or len(set(rows[0])) != actual or any(not x.strip() for x in rows[0]):
            raise ValueError(f"{name}: invalid schema ({actual} fields; expected {expected})")
        counts[name] = actual
    return counts


def validate_edition(edition: dict, tool_counts: dict[str, int]) -> dict:
    md = edition["md"]
    text = md.read_text(encoding="utf-8")
    if re.search(r"\b(TODO|TBD|FIXME|PLACEHOLDER)\b", text, re.I):
        raise ValueError(f"Placeholder found in {md}")
    docx = build_docx(md, edition["title"], edition["status"])
    with zipfile.ZipFile(docx) as archive:
        if "word/document.xml" not in archive.namelist():
            raise ValueError(f"DOCX structure invalid: {docx}")
    pdf = convert_pdf(docx)
    reader = PdfReader(str(pdf))
    pages = len(reader.pages)
    extracted = "\n".join((page.extract_text() or "") for page in reader.pages)
    words = len(re.findall(r"\b[\wÀ-ÿ][\wÀ-ÿ'-]*\b", extracted))
    if pages < 2 or words < 700:
        raise ValueError(f"Searchable PDF validation failed: {pdf} ({pages} pages, {words} words)")
    report = {
        "status": "PASS",
        "locale": edition["locale"],
        "markdown_h1": len(re.findall(r"^# ", text, re.M)),
        "markdown_h2": len(re.findall(r"^## ", text, re.M)),
        "docx_zip_integrity": "PASS",
        "pdf_searchable_text": "PASS",
        "pdf_pages": pages,
        "pdf_extracted_words": words,
        "tool_field_counts": tool_counts,
        "human_review_limitations": [
            "Native-language editorial approval is not represented as completed.",
            "Organization-specific legal, regulatory, accessibility, certification, and audit review is not represented as completed.",
            "Automated validation does not prove program or control effectiveness.",
        ],
    }
    qa_dir = edition["qa"]
    if edition["locale"] == "en":
        (qa_dir / "ENGLISH_PACKAGE_QA.json").write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
        md_name = "ENGLISH_PACKAGE_QA.md"
        sums_name = "ENGLISH_SHA256SUMS.txt"
    else:
        (qa_dir / "LOCALIZED_QA.json").write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
        md_name = "LOCALIZED_QA.md"
        sums_name = "SHA256SUMS.txt"
    qa_lines = [
        f"# Package QA - {edition['locale']}", "", "- Automated result: **PASS**",
        f"- H1 headings: {report['markdown_h1']}", f"- H2 headings: {report['markdown_h2']}",
        "- DOCX ZIP integrity: PASS", "- PDF searchable text: PASS",
        f"- PDF pages: {pages}", f"- Extracted PDF words: {words}", "",
        "## Human-review limitations", "",
        "Native-language editorial approval, organization-specific legal or regulatory review, accessibility and assistive-technology testing, standards certification, formal audit assurance, and approval of organization-specific decision rights are not represented as completed.", "",
    ]
    (qa_dir / md_name).write_text("\n".join(qa_lines), encoding="utf-8")
    checksum_paths = [md, docx, pdf]
    if edition["locale"] == "en":
        checksum_paths.extend(ROOT / "tools" / name for name in TOOLS)
    (qa_dir / sums_name).write_text(
        "".join(f"{sha256(path)}  {Path(os.path.relpath(path, qa_dir)).as_posix()}\n" for path in checksum_paths),
        encoding="utf-8",
    )
    return report


def main() -> None:
    tool_counts = validate_tools()
    reports = [validate_edition(edition, tool_counts) for edition in EDITIONS]
    english = reports[0]
    for report in reports[1:]:
        if report["markdown_h1"] != english["markdown_h1"] or report["markdown_h2"] != english["markdown_h2"]:
            raise ValueError(f"Heading parity failed for {report['locale']}")
    assembly = {
        "status": "PASS", "repository_release": "1.11", "edition_count": 3,
        "tool_field_counts": tool_counts, "editions": reports,
    }
    out = ROOT / "English"
    (out / "ENGLISH_ASSEMBLY_REPORT.json").write_text(json.dumps(assembly, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    (out / "ENGLISH_ASSEMBLY_REPORT.md").write_text(
        "# English Assembly Report\n\n- Status: **PASS**\n- Repository release: 1.11\n- Editions: 3\n- CSV tools: 6\n- Heading parity: PASS\n\nAutomated results do not replace qualified human review.\n",
        encoding="utf-8",
    )
    print(json.dumps(assembly, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"FAIL: {exc}", file=sys.stderr)
        raise
