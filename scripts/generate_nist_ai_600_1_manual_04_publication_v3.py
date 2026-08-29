#!/usr/bin/env python3
"""Manual 04 publication generator v3: localize generated labels and prevent linked-footer duplication."""
from __future__ import annotations

from docx import Document

import generate_nist_ai_600_1_manual_04_publication_v2 as v2

core = v2.core
_adapter_build_docx = core.build_docx
_v2_render = core.render_mermaid_memory_graphic
_v2_alt = core.set_image_alt_text
_original_add_footer = v2.adapter.add_footer

LOCALIZED = {
    "es-419": {
        "implementation_heading": "Rutas de implementación y modelo operativo",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "control_prefix": "Revisión de fuente controlada:",
        "language_label": "Idioma:",
        "control_tail": "Perfil de IA Generativa NIST AI 600-1 con el marco principal AI RMF 1.0; control de versiones y vigilancia de fuentes.",
        "boundary": "Límite de aseguramiento: Este manual voluntario de implementación no certifica un sistema de IA generativa, no establece cumplimiento legal, no demuestra por sí solo una IA confiable ni proporciona una opinión de auditoría; no todas las acciones sugeridas aplican a cada actor, sistema, caso de uso o contexto del ciclo de vida.",
        "figure_prefix": "Figura",
        "caption": "Gráfico de memoria de implementación",
        "graphic_title": "Gráfico de memoria del Manual 04",
        "footer_status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACION",
    },
    "pt-BR": {
        "implementation_heading": "Caminhos de implementação e modelo operacional",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "control_prefix": "Revisão da fonte controlada:",
        "language_label": "Idioma:",
        "control_tail": "Perfil de IA Generativa NIST AI 600-1 com o framework principal AI RMF 1.0; controle de versão e monitoramento de fontes.",
        "boundary": "Limite de asseguração: Este manual voluntário de implementação não certifica um sistema de IA generativa, não estabelece conformidade legal, não demonstra por si só uma IA confiável nem fornece uma opinião de auditoria; nem toda ação sugerida se aplica a cada ator, sistema, caso de uso ou contexto do ciclo de vida.",
        "figure_prefix": "Figura",
        "caption": "Gráfico de memória de implementação",
        "graphic_title": "Gráfico de memória do Manual 04",
        "footer_status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACAO",
    },
}


def language_from_path(path) -> str:
    s = str(path)
    if "es_419" in s or "ES-419" in s:
        return "es-419"
    if "pt_BR" in s or "PT-BR" in s:
        return "pt-BR"
    return "en"


def add_footer_once(section, language: str):
    """Prevent duplicate text when Word sections share the same footer part."""
    existing = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
    if "Manual 04 |" in existing:
        return
    _original_add_footer(section, language)


def render_mermaid_memory_graphic(block: str, out_path, title: str) -> str:
    lang = language_from_path(out_path)
    if lang in LOCALIZED:
        number = title.rsplit(" ", 1)[-1]
        title = f"{LOCALIZED[lang]['graphic_title']} {number}"
    return _v2_render(block, out_path, title)


def set_image_alt_text(inline_shape, title: str, description: str):
    if "Gráfico de memoria del Manual 04" in description:
        title = title.replace("Manual 04 memory graphic", "Gráfico de memoria del Manual 04")
    elif "Gráfico de memória do Manual 04" in description:
        title = title.replace("Manual 04 memory graphic", "Gráfico de memória do Manual 04")
    return _v2_alt(inline_shape, title, description)


def replace_paragraph_text(paragraph, text: str):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def localize_footer(doc: Document, language: str):
    status = (
        LOCALIZED[language]["footer_status"]
        if language in LOCALIZED
        else "CONTROLLED PUBLICATION QA CANDIDATE"
    )
    seen = set()
    for section in doc.sections:
        footer = section.footer
        key = id(footer._element)
        if key in seen:
            continue
        seen.add(key)
        for paragraph in footer.paragraphs:
            if "Manual 04 |" not in paragraph.text:
                continue
            if not paragraph.runs:
                continue
            paragraph.runs[0].text = f"Manual 04 | {status} | "
            for run in paragraph.runs[1:]:
                run.text = ""
            core.set_run_font(paragraph.runs[0], language, size=8)


def build_docx(source, out_path, image_dir, source_head: str):
    graphic_count = _adapter_build_docx(source, out_path, image_dir, source_head)
    language = source.language
    doc = Document(out_path)

    if language in LOCALIZED:
        loc = LOCALIZED[language]
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text == "Implementation paths and operating model":
                replace_paragraph_text(paragraph, loc["implementation_heading"])
            elif text == "Controlled 32-chapter manual":
                replace_paragraph_text(paragraph, loc["chapter_heading"])
            elif text.startswith("Controlled source revision:"):
                replace_paragraph_text(
                    paragraph,
                    f"{loc['control_prefix']} {source_head} | {loc['language_label']} {language} | {loc['control_tail']}",
                )
            elif text.startswith("Assurance boundary:"):
                replace_paragraph_text(paragraph, loc["boundary"])
            elif text.startswith("Figure ") and text.endswith("Implementation memory graphic"):
                number = text.split()[1].rstrip(".")
                replace_paragraph_text(paragraph, f"{loc['figure_prefix']} {number}. {loc['caption']}")

    localize_footer(doc, language)

    # Re-apply language/font metadata after replacing generated strings.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            core.set_run_font(run, language)
    doc.save(out_path)
    return graphic_count


# The adapter build function performs global lookup of its own add_footer name.
v2.adapter.add_footer = add_footer_once
core.add_footer = add_footer_once
core.render_mermaid_memory_graphic = render_mermaid_memory_graphic
core.set_image_alt_text = set_image_alt_text
core.build_docx = build_docx

if __name__ == "__main__":
    raise SystemExit(core.main())
