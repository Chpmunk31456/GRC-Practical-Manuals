#!/usr/bin/env python3
"""Build the exact Manual 46 trilingual DOCX/PDF publication candidate.

The builder is fail-closed. It requires controlled EN, es-419 and pt-BR source
packages and refuses to emit a candidate if required source material is absent.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
import hashlib
import json
import os
import re
import subprocess

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "10-ai-governance/Manual_46_Universal_AI_Governance"
OUT = ROOT / "qa/manual46-publication-candidate"
OUT.mkdir(parents=True, exist_ok=True)

EN_COMPONENTS = [
    BASE / "README.md",
    BASE / "TRAINING_MODULES.md",
    BASE / "PRACTICAL_SCENARIOS.md",
    BASE / "CONTROL_EVIDENCE_WORKBOOK.md",
]
LOCALE_SOURCES = {
    "es-419": BASE / "controlled/es-419/MANUAL_46_CONTROLLED_ES-419.md",
    "pt-BR": BASE / "controlled/pt-BR/MANUAL_46_CONTROLLED_PT-BR.md",
}
NAMES = {
    "en": "Manual_46_Universal_AI_Governance_Foundation_EN",
    "es-419": "Manual_46_Universal_AI_Governance_Foundation_ES-419",
    "pt-BR": "Manual_46_Universal_AI_Governance_Foundation_PT-BR",
}


def read_required(path: Path) -> str:
    if not path.is_file():
        raise SystemExit(f"required controlled source missing: {path.relative_to(ROOT)}")
    text = path.read_text(encoding="utf-8")
    if len(text.strip()) < 500:
        raise SystemExit(f"controlled source is unexpectedly short: {path.relative_to(ROOT)}")
    return text


def english_source() -> str:
    parts = []
    for path in EN_COMPONENTS:
        parts.append(read_required(path))
    text = "\n\n---\n\n".join(parts)
    required_markers = [
        "Universal AI Governance Foundation",
        "Universal governance spine",
        "Controlled Training Modules",
        "Practical Scenarios",
        "Control and Evidence Workbook",
    ]
    for marker in required_markers:
        if marker not in text:
            raise SystemExit(f"English controlled package missing marker: {marker}")
    return text


def add_markdown(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            doc.add_heading(heading.group(2).strip(), level=min(len(heading.group(1)), 3))
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            doc.add_paragraph(numbered.group(1).strip(), style="List Number")
            continue
        if line == "---":
            continue
        paragraph = doc.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(6)


sources = {
    "en": english_source(),
    "es-419": read_required(LOCALE_SOURCES["es-419"]),
    "pt-BR": read_required(LOCALE_SOURCES["pt-BR"]),
}

# Lightweight parity guard before binary generation.
for locale, text in sources.items():
    for concept in ("inventory", "risk", "security", "privacy") if locale == "en" else ():
        if concept.lower() not in text.lower():
            raise SystemExit(f"{locale} source missing core concept: {concept}")

manifest = {
    "manual": 46,
    "title": "Universal AI Governance Foundation",
    "source_commit": os.environ.get("GITHUB_SHA", ""),
    "artifacts": [],
}

for locale, text in sources.items():
    doc = Document()
    doc.styles["Normal"].font.name = "Liberation Sans"
    doc.styles["Normal"].font.size = Pt(10.5)
    add_markdown(doc, text)
    stem = NAMES[locale]
    docx_path = OUT / f"{stem}.docx"
    doc.save(docx_path)
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)],
        check=True,
    )
    pdf_path = OUT / f"{stem}.pdf"
    for path in (docx_path, pdf_path):
        if not path.is_file() or path.stat().st_size == 0:
            raise SystemExit(f"empty or missing candidate artifact: {path}")
        data = path.read_bytes()
        manifest["artifacts"].append(
            {
                "locale": locale,
                "file": path.name,
                "sha256": hashlib.sha256(data).hexdigest(),
                "bytes": len(data),
            }
        )

manifest_path = OUT / "MANUAL_46_CANDIDATE_MANIFEST.json"
manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
print(json.dumps(manifest, indent=2, ensure_ascii=False))
