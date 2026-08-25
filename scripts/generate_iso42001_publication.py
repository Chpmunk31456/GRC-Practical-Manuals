#!/usr/bin/env python3
"""Build controlled Manual 02 publication-QA artifacts.

The English Markdown file remains the semantic source of truth. Localized
masters are deterministic, mechanically consolidated derivatives of the four
review parts recorded in the Manual 02 baseline. Localized outputs always carry
the human-review-required status and are never represented as final or
ISO-authorized translations.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "02-management-systems" / "ISO_IEC_42001_AIMS"
BASELINE = ROOT / ".compliance" / "iso-42001-manual-02-baseline.json"
REFERENCE_DOCX = MANUAL / "English" / "ISO_IEC_42001_2023_Practical_AIMS_Manual_English_v1.0.docx"
OUTPUT_DIR = MANUAL / "publication" / "qa-candidate"

FILENAMES = {
    "en": "Manual_02_ISO_IEC_42001_AI_Management_System_EN",
    "es-419": "Manual_02_ISO_IEC_42001_AI_Management_System_ES-419",
    "pt-BR": "Manual_02_ISO_IEC_42001_AI_Management_System_PT-BR",
}

TITLES = {
    "en": "Manual 02 — ISO/IEC 42001 AI Management System Implementation",
    "es-419": "Manual 02 — Implementación del Sistema de Gestión de IA ISO/IEC 42001",
    "pt-BR": "Manual 02 — Implementação do Sistema de Gestão de IA ISO/IEC 42001",
}

STATUS = {
    "en": "PUBLICATION QA CANDIDATE — CONTROLLED ENGLISH SOURCE",
    "es-419": "DRAFT — HUMAN SEMANTIC REVIEW REQUIRED",
    "pt-BR": "DRAFT — HUMAN SEMANTIC REVIEW REQUIRED",
}

LANG_WORD = {"en": "en-US", "es-419": "es-419", "pt-BR": "pt-BR"}

LOCALIZED_CONTROL = {
    "es-419": (
        "Control de publicación: Este documento es un borrador derivado mecánicamente de los cuatro "
        "archivos fuente localizados de la rama indicada. Requiere revisión semántica humana. No es una "
        "traducción autorizada por ISO y su generación no demuestra certificación, conformidad, cumplimiento "
        "legal ni aseguramiento de auditoría. La asistencia de IA se utilizó conforme a la divulgación del "
        "repositorio; la autoría y responsabilidad humana permanecen con Alberto (Al) Leiva."
    ),
    "pt-BR": (
        "Controle de publicação: Este documento é um rascunho derivado mecanicamente dos quatro arquivos "
        "fonte localizados da ramificação indicada. Requer revisão semântica humana. Não é uma tradução "
        "autorizada pela ISO, e sua geração não demonstra certificação, conformidade, cumprimento legal nem "
        "asseguração de auditoria. A assistência de IA foi utilizada conforme a divulgação do repositório; "
        "a autoria e a responsabilidade humana permanecem com Alberto (Al) Leiva."
    ),
}

IMAGE_RE = re.compile(
    r'<img\s+src="([^"]+)"\s+style="width:([^;]+);height:([^"]+)"\s+alt="([^"]+)"\s*/?>'
)

ENGLISH_TEXT_REPLACEMENTS = {
    "Current-information note: Verified July 14, 2026. ISO/IEC 42001:2023 remains the published AIMS requirements standard. ISO/IEC 42005:2025 provides AI system impact-assessment guidance. ISO/IEC 42006:2025 adds requirements for bodies auditing and certifying AIMS. ISO 19011:2026 is the current management-system audit guideline. ISO/IEC 42003 and 42007 remain under development and are not treated as requirements here.":
        "Current-information note: Verified August 24, 2026. ISO/IEC 42001:2023 remains the published AIMS requirements standard. ISO/IEC 42005:2025 provides AI system impact-assessment guidance. ISO/IEC 42006:2025 adds requirements for bodies auditing and certifying AIMS. ISO 19011:2026 is the current management-system audit guideline. ISO/IEC 42003 remains an approved work item and ISO/IEC 42007 has advanced to draft international standard status; both remain under development and are not treated as requirements here.",
    "Keep one source of truth and map it to ISO 27001, ISO 9001, privacy, legal, NIST AI RMF, and sector obligations rather than duplicating records.":
        "Keep one source of truth and map it to ISO/IEC 27001:2022, ISO 9001, privacy, legal, NIST AI RMF, and sector obligations rather than duplicating records.",
    "Annex A.4 requires visibility into the data, tools, system/compute, and people needed across the AI life cycle.":
        "When Annex A.4 controls are selected through the organization's risk-treatment process and Statement of Applicability, implementation should maintain visibility into the data, tools, systems and computing resources, and people needed across the AI life cycle.",
    "Annex A.7 requires governed data acquisition, quality, provenance, and preparation for AI development, enhancement, and operation.":
        "When Annex A.7 controls are selected through the organization's risk-treatment process and Statement of Applicability, implementation should govern data acquisition, quality, provenance, and preparation for AI development, enhancement, and operation.",
    "Annex A.8 requires useful information for users and interested parties, plus reporting and incident communication.":
        "When Annex A.8 controls are selected through the organization's risk-treatment process and Statement of Applicability, implementation should provide useful information for users and interested parties, together with reporting and incident communication.",
}

ENGLISH_HYPERLINK_REPLACEMENTS = {
    "https://github.com/Giskard-AI/giskard": "https://github.com/Giskard-AI/giskard-oss",
    "https://github.com/Azure/PyRIT": "https://github.com/microsoft/PyRIT",
    "https://microsoft.github.io/presidio/": "https://presidio.dataprivacystack.org/",
}

ENGLISH_REFERENCE_LINKS = [
    ("ISO/IEC 27001:2022 information-security management-system requirements", "https://www.iso.org/standard/27001"),
    ("ISO/IEC 27001:2022/Amd 1:2024 climate-action changes", "https://www.iso.org/standard/88435.html"),
    ("ISO/IEC 17021-1:2015 management-system certification bodies", "https://www.iso.org/standard/61651.html"),
]


def xml_attr(element: OxmlElement, name: str, value: str) -> None:
    element.set(qn(name), value)


def add_hyperlink(paragraph, display_text: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    xml_attr(hyperlink, "r:id", relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    xml_attr(color, "w:val", "0563C1")
    underline = OxmlElement("w:u")
    xml_attr(underline, "w:val", "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = display_text
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def synchronize_english_controlled_source(document: Document) -> None:
    """Apply reviewed Markdown changes to the preserved English DOCX source."""
    for old, new in ENGLISH_TEXT_REPLACEMENTS.items():
        matches = [paragraph for paragraph in document.paragraphs if old in paragraph.text]
        if len(matches) != 1:
            raise ValueError(f"English DOCX source-sync marker matched {len(matches)} times: {old[:60]}")
        matches[0].text = matches[0].text.replace(old, new)

    for relationship in document.part.rels.values():
        replacement = ENGLISH_HYPERLINK_REPLACEMENTS.get(relationship.target_ref)
        if replacement:
            relationship._target = replacement
    targets = {relationship.target_ref for relationship in document.part.rels.values()}
    missing_targets = set(ENGLISH_HYPERLINK_REPLACEMENTS.values()) - targets
    if missing_targets:
        raise ValueError(f"English DOCX hyperlink synchronization failed: {sorted(missing_targets)}")

    existing_text = {paragraph.text.strip() for paragraph in document.paragraphs}
    anchor = next(
        (paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "ISO/IEC JTC 1/SC 42 catalogue"),
        None,
    )
    if anchor is None:
        raise ValueError("English DOCX reference-list anchor not found")
    for display_text, target in ENGLISH_REFERENCE_LINKS:
        if display_text in existing_text:
            continue
        paragraph = anchor.insert_paragraph_before(style="List Bullet")
        add_hyperlink(paragraph, display_text, target)


def run_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    xml_attr(begin, "w:fldCharType", "begin")
    instr = OxmlElement("w:instrText")
    xml_attr(instr, "xml:space", "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    xml_attr(separate, "w:fldCharType", "separate")
    end = OxmlElement("w:fldChar")
    xml_attr(end, "w:fldCharType", "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, end])


def add_callout_format(paragraph_xml: OxmlElement, index: int, total: int) -> None:
    props = paragraph_xml.get_or_add_pPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    xml_attr(shading, "w:val", "clear")
    xml_attr(shading, "w:color", "auto")
    xml_attr(shading, "w:fill", "FFF4DA")

    borders = props.find(qn("w:pBdr"))
    if borders is not None:
        props.remove(borders)
    borders = OxmlElement("w:pBdr")
    sides = ["left", "right"]
    if index == 0:
        sides.append("top")
    if index == total - 1:
        sides.append("bottom")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        xml_attr(border, "w:val", "single")
        xml_attr(border, "w:sz", "8")
        xml_attr(border, "w:space", "4")
        xml_attr(border, "w:color", "D5B46B")
        borders.append(border)
    props.append(borders)

    indent = props.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        props.append(indent)
    xml_attr(indent, "w:left", "120")
    xml_attr(indent, "w:right", "120")


def remove_layout_tables(document: Document) -> None:
    callouts = [table for table in document.tables if len(table.rows) == 1 and len(table.columns) == 1]
    for table in callouts:
        table_xml = table._tbl
        cell_paragraphs = list(table.cell(0, 0)._tc.p_lst)
        for index, paragraph in enumerate(cell_paragraphs):
            copied = deepcopy(paragraph)
            add_callout_format(copied, index, len(cell_paragraphs))
            table_xml.addprevious(copied)
        table_xml.getparent().remove(table_xml)


def set_keep_next(paragraph) -> None:
    props = paragraph._p.get_or_add_pPr()
    if props.find(qn("w:keepNext")) is None:
        props.append(OxmlElement("w:keepNext"))


def apply_docx_controls(
    path: Path,
    language: str,
    source_commit: str,
    source_branch: str,
    generation_date: str,
    alt_texts: list[str],
) -> None:
    document = Document(path)
    if language == "en":
        synchronize_english_controlled_source(document)
    remove_layout_tables(document)

    props = document.core_properties
    props.title = TITLES[language]
    props.subject = (
        f"Manual 02 controlled publication-QA artifact; {STATUS[language]}; "
        f"source {source_branch}@{source_commit}"
    )
    props.author = "Alberto (Al) Leiva"
    props.keywords = "ISO/IEC 42001, AI management system, AIMS, GRC, implementation, audit evidence"
    props.comments = (
        "Document generation and repository QA do not establish certification, conformity, legal "
        "compliance, or audit assurance."
    )

    if language == "en":
        nonempty = [p for p in document.paragraphs if p.text.strip()]
        if nonempty:
            first = nonempty[0]
            first.text = TITLES[language]
            first.style = document.styles["Title"]
        preface = next((p for p in document.paragraphs if p.text.strip() == "Preface"), None)
        if preface is not None:
            control = preface.insert_paragraph_before(
                f"Publication control — Version: First Edition / v1.0; Language: EN; "
                f"Source: {source_branch}@{source_commit}; Generated: {generation_date}; "
                f"Review status: {STATUS[language]}. AI assistance is disclosed in the repository; "
                "human authorship and accountability remain with Alberto (Al) Leiva. Document "
                "generation alone does not establish certification, conformity, legal compliance, "
                "or audit assurance."
            )
            control.style = document.styles["Normal"]

    for style in document.styles:
        element = style._element
        rpr = element.get_or_add_rPr()
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        xml_attr(lang, "w:val", LANG_WORD[language])
        xml_attr(lang, "w:eastAsia", LANG_WORD[language])

    for paragraph in document.paragraphs:
        name = paragraph.style.name or ""
        if name.startswith("Heading") or name in {"Title", "Subtitle"}:
            set_keep_next(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if re.match(r"^(Figure|Figura)\s+\d+\.", paragraph.text.strip()):
            try:
                paragraph.style = document.styles["Caption"]
            except KeyError:
                pass
            set_keep_next(paragraph)

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = language != "en"
        if language != "en":
            table_properties = table._tbl.tblPr
            existing_borders = table_properties.find(qn("w:tblBorders"))
            if existing_borders is not None:
                table_properties.remove(existing_borders)
            borders = OxmlElement("w:tblBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                xml_attr(border, "w:val", "single")
                xml_attr(border, "w:sz", "4")
                xml_attr(border, "w:color", "808080")
                borders.append(border)
            table_properties.append(borders)

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        if language != "en":
                            run.font.size = Pt(9.5)

        if table.rows:
            row_props = table.rows[0]._tr.get_or_add_trPr()
            if row_props.find(qn("w:tblHeader")) is None:
                row_props.append(OxmlElement("w:tblHeader"))
            for cell in table.rows[0].cells:
                cell_properties = cell._tc.get_or_add_tcPr()
                shading = cell_properties.find(qn("w:shd"))
                if shading is None:
                    shading = OxmlElement("w:shd")
                    cell_properties.append(shading)
                xml_attr(shading, "w:fill", "1F4E78")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = None
                        color = run._r.get_or_add_rPr().find(qn("w:color"))
                        if color is None:
                            color = OxmlElement("w:color")
                            run._r.get_or_add_rPr().append(color)
                        xml_attr(color, "w:val", "FFFFFF")
        for row in table.rows:
            row_props = row._tr.get_or_add_trPr()
            cant_split = row_props.find(qn("w:cantSplit"))
            if cant_split is None:
                row_props.append(OxmlElement("w:cantSplit"))

    english_section = Document(REFERENCE_DOCX).sections[0]
    for section in document.sections:
        if language != "en":
            section.page_width = english_section.page_width
            section.page_height = english_section.page_height
            section.top_margin = english_section.top_margin
            section.bottom_margin = english_section.bottom_margin
            section.left_margin = english_section.left_margin
            section.right_margin = english_section.right_margin
            section.header_distance = english_section.header_distance
            section.footer_distance = english_section.footer_distance
        section.header.is_linked_to_previous = False
        header = section.header.paragraphs[0]
        header.text = f"Manual 02 | {language} | First Edition / v1.0 | {STATUS[language]}"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        section.footer.is_linked_to_previous = False
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(f"{source_branch}@{source_commit[:12]} | {generation_date} | Page ")
        run_field(footer, "PAGE")
        footer.add_run(" of ")
        run_field(footer, "NUMPAGES")

    doc_prs = document._element.xpath(".//wp:docPr")
    if len(doc_prs) != len(alt_texts):
        raise ValueError(
            f"{language}: generated DOCX has {len(doc_prs)} images; expected {len(alt_texts)}"
        )
    for index, (doc_pr, alt_text) in enumerate(zip(doc_prs, alt_texts, strict=True), start=1):
        doc_pr.set("name", f"Manual 02 Figure {index}")
        doc_pr.set("title", f"Figure {index}")
        doc_pr.set("descr", alt_text)

    document.save(path)


def normalize_images(text: str) -> tuple[str, list[str]]:
    alt_texts: list[str] = []

    def replacement(match: re.Match[str]) -> str:
        source, width, height, alt = match.groups()
        alt_texts.append(alt)
        return f"![{alt}]({source}){{width={width} height={height}}}"

    return IMAGE_RE.sub(replacement, text), alt_texts


def localized_master(
    language: str,
    parts: list[Path],
    source_commit: str,
    source_branch: str,
    generation_date: str,
) -> tuple[str, list[str]]:
    source_parts = []
    for path in parts:
        mechanically_normalized = "\n".join(
            line.rstrip() for line in path.read_text(encoding="utf-8").splitlines()
        )
        source_parts.append(mechanically_normalized.rstrip())
    combined = "\n\n".join(source_parts) + "\n"
    marker = "| **Qué hace este manual:**" if language == "es-419" else "| **O que este manual faz:**"
    start = combined.find(marker)
    if start < 0:
        raise ValueError(f"{language}: controlled opening table marker not found")

    yaml = (
        "---\n"
        f'title: "{TITLES[language]}"\n'
        f'subtitle: "{STATUS[language]}"\n'
        'author: "Alberto (Al) Leiva"\n'
        f'date: "{generation_date}"\n'
        f'lang: "{language}"\n'
        'version: "First Edition / v1.0"\n'
        f'source-branch: "{source_branch}"\n'
        f'source-commit: "{source_commit}"\n'
        f'review-status: "{STATUS[language]}"\n'
        "---\n\n"
        f"> **{LOCALIZED_CONTROL[language]}**\n\n"
        f"> **Source control:** `{source_branch}` @ `{source_commit}` · {generation_date} · First Edition / v1.0\n\n"
    )
    normalized, alt_texts = normalize_images(yaml + combined[start:])
    if len(alt_texts) != 10:
        raise ValueError(f"{language}: consolidated master has {len(alt_texts)} images; expected 10")
    return normalized, alt_texts


def run_pandoc_odt(master: Path, output: Path, language: str) -> None:
    resource_path = ":".join(
        [str(master.parent), str(MANUAL), str(MANUAL / "assets"), str(ROOT)]
    )
    publication_input = re.sub(
        rf"(assets/{re.escape(language)}/media/image\d+)\.png",
        r"\1.svg",
        master.read_text(encoding="utf-8"),
    )
    subprocess.run(
        [
            "pandoc",
            "-",
            "--from=markdown+yaml_metadata_block+raw_html+pipe_tables+link_attributes",
            "--to=odt",
            "--standalone",
            "--toc",
            "--toc-depth=3",
            f"--resource-path={resource_path}",
            "-o",
            str(output),
        ],
        cwd=master.parent,
        input=publication_input,
        text=True,
        check=True,
    )


def convert_odt_docx(odt_path: Path, docx_path: Path) -> None:
    with (
        tempfile.TemporaryDirectory(prefix="manual02-odt-profile-") as profile,
        tempfile.TemporaryDirectory(prefix="manual02-odt-output-") as output,
        tempfile.TemporaryDirectory(prefix="manual02-odt-input-") as input_dir,
    ):
        conversion_input = Path(input_dir) / odt_path.name
        shutil.copyfile(odt_path, conversion_input)
        subprocess.run(
            [
                "soffice",
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "docx:Office Open XML Text",
                "--outdir",
                output,
                str(conversion_input),
            ],
            check=True,
        )
        generated = Path(output) / f"{odt_path.stem}.docx"
        if not generated.is_file() or generated.stat().st_size == 0:
            raise ValueError(f"ODT conversion did not produce {generated}")
        shutil.copyfile(generated, docx_path)


def convert_pdf(docx_path: Path, pdf_path: Path) -> None:
    if pdf_path.exists():
        pdf_path.unlink()
    with (
        tempfile.TemporaryDirectory(prefix="manual02-lo-profile-") as profile,
        tempfile.TemporaryDirectory(prefix="manual02-lo-output-") as output,
        tempfile.TemporaryDirectory(prefix="manual02-lo-input-") as input_dir,
    ):
        conversion_input = Path(input_dir) / docx_path.name
        shutil.copyfile(docx_path, conversion_input)
        subprocess.run(
            [
                "soffice",
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                output,
                str(conversion_input),
            ],
            check=True,
        )
        generated = Path(output) / f"{docx_path.stem}.pdf"
        if not generated.is_file() or generated.stat().st_size == 0:
            raise ValueError(f"PDF conversion did not produce {generated}")
        optimized = Path(output) / f"{docx_path.stem}-optimized.pdf"
        subprocess.run(
            [
                "gs",
                "-q",
                "-sDEVICE=pdfwrite",
                "-dCompatibilityLevel=1.7",
                "-dPDFSETTINGS=/ebook",
                "-dNOPAUSE",
                "-dBATCH",
                f"-sOutputFile={optimized}",
                str(generated),
            ],
            check=True,
        )
        if not optimized.is_file() or optimized.stat().st_size == 0:
            raise ValueError(f"PDF optimization did not produce {optimized}")
        shutil.copyfile(optimized, pdf_path)
    if not pdf_path.is_file() or pdf_path.stat().st_size == 0:
        raise ValueError(f"PDF conversion did not produce {pdf_path}")


def refuse_existing(paths: list[Path], force: bool) -> None:
    existing = [path for path in paths if path.exists()]
    if existing and not force:
        listing = "\n".join(f"  {path.relative_to(ROOT)}" for path in existing)
        raise FileExistsError(f"refusing to overwrite generated output without --force:\n{listing}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-commit", required=True)
    parser.add_argument("--source-branch", default="build/iso-iec-42001-manual-02-2026")
    parser.add_argument("--generation-date", required=True)
    parser.add_argument("--force", action="store_true")
    args = parser.parse_args()

    baseline = json.loads(BASELINE.read_text(encoding="utf-8"))
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    master_paths = {
        "es-419": MANUAL / "translations" / "es-419" / "source" / f"{FILENAMES['es-419']}.md",
        "pt-BR": MANUAL / "translations" / "pt-BR" / "source" / f"{FILENAMES['pt-BR']}.md",
    }
    docx_paths = {lang: OUTPUT_DIR / f"{stem}.docx" for lang, stem in FILENAMES.items()}
    pdf_paths = {lang: OUTPUT_DIR / f"{stem}.pdf" for lang, stem in FILENAMES.items()}
    refuse_existing(list(master_paths.values()) + list(docx_paths.values()) + list(pdf_paths.values()), args.force)

    alt_by_language: dict[str, list[str]] = {}
    for language in ("es-419", "pt-BR"):
        parts = [MANUAL / relative for relative in baseline["localized_full_source_parts"][language]]
        text, alt_texts = localized_master(
            language, parts, args.source_commit, args.source_branch, args.generation_date
        )
        master_paths[language].write_text(text, encoding="utf-8", newline="\n")
        alt_by_language[language] = alt_texts

    english_markdown = MANUAL / baseline["english_markdown_master"]
    _, english_alt = normalize_images(english_markdown.read_text(encoding="utf-8"))
    if len(english_alt) != 10:
        raise ValueError(f"en: controlled Markdown has {len(english_alt)} images; expected 10")
    alt_by_language["en"] = english_alt

    shutil.copyfile(REFERENCE_DOCX, docx_paths["en"])
    apply_docx_controls(
        docx_paths["en"],
        "en",
        args.source_commit,
        args.source_branch,
        args.generation_date,
        alt_by_language["en"],
    )

    with tempfile.TemporaryDirectory(prefix="manual02-odt-stage-") as odt_stage:
        for language in ("es-419", "pt-BR"):
            odt_path = Path(odt_stage) / f"{FILENAMES[language]}.odt"
            run_pandoc_odt(master_paths[language], odt_path, language)
            convert_odt_docx(odt_path, docx_paths[language])
            apply_docx_controls(
                docx_paths[language],
                language,
                args.source_commit,
                args.source_branch,
                args.generation_date,
                alt_by_language[language],
            )

    for language in ("en", "es-419", "pt-BR"):
        convert_pdf(docx_paths[language], pdf_paths[language])

    for language, path in docx_paths.items():
        generated = Document(path)
        if len(generated.tables) != 33 or len(generated.inline_shapes) != 10:
            raise ValueError(
                f"{language}: post-conversion DOCX integrity changed unexpectedly "
                f"({len(generated.tables)} tables, {len(generated.inline_shapes)} images)"
            )
        if language != "en" and any(
            len(table.rows) == 1 and len(table.columns) == 1 for table in generated.tables
        ):
            raise ValueError(f"{language}: a layout-only table reappeared after conversion")

    for path in list(master_paths.values()) + list(docx_paths.values()) + list(pdf_paths.values()):
        if not path.is_file() or path.stat().st_size == 0:
            raise ValueError(f"generated output missing or empty: {path}")
        print(path.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
