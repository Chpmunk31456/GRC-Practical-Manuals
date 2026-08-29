#!/usr/bin/env python3
"""Manual 05 publication generator v3.

Preserves v2 source-relative completeness validation, makes footer creation
idempotent across Word sections, and localizes generator-owned captions and
control-line suffixes in the es-419 / pt-BR publication candidates.
"""
from __future__ import annotations

from docx import Document

import generate_ai_auditing_assurance_manual_05_publication as v1
import generate_ai_auditing_assurance_manual_05_publication_v2 as v2  # noqa: F401

_original_add_footer = v1.add_footer
_original_build_docx = v1.build_docx

CONTROL_TAIL = {
    "en": "controlled assurance baseline and source watch retained.",
    "es-419": "se conservan la línea base de aseguramiento controlada y la vigilancia de fuentes.",
    "pt-BR": "a linha de base de asseguração controlada e o monitoramento de fontes são mantidos.",
}


def add_footer(section, language: str):
    existing = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
    if "Manual 05" in existing:
        return
    _original_add_footer(section, language)


def replace_paragraph_text(paragraph, text: str, language: str):
    if not paragraph.runs:
        paragraph.add_run(text)
    else:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    for run in paragraph.runs:
        v1.core.set_run_font(run, language)


def build_docx(source, out_path, image_dir, source_head: str):
    count = _original_build_docx(source, out_path, image_dir, source_head)
    language = source.language
    meta = v1.LANG_META[language]
    doc = Document(out_path)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(meta["control_prefix"]):
            replace_paragraph_text(
                paragraph,
                f"{meta['control_prefix']} {source_head} | {meta['language_label']} {language} | {CONTROL_TAIL[language]}",
                language,
            )
            continue

        if language in {"es-419", "pt-BR"} and text.startswith("Figure ") and "Implementation memory graphic" in text:
            number = text.split(".", 1)[0].split()[-1]
            replace_paragraph_text(
                paragraph,
                f"Figura {number}. {meta['figure_caption']}",
                language,
            )

    # Localize generator-owned document metadata where the edition is localized.
    if language == "es-419":
        doc.core_properties.subject = f"Artefacto controlado de QA de publicación del Manual 05; idioma es-419; revisión de fuente {source_head}"
        doc.core_properties.comments = "La generación y el QA no establecen independencia, competencia, cumplimiento legal, certificación, conformidad ni una opinión de auditoría."
    elif language == "pt-BR":
        doc.core_properties.subject = f"Artefato controlado de QA de publicação do Manual 05; idioma pt-BR; revisão da fonte {source_head}"
        doc.core_properties.comments = "A geração e o QA não estabelecem independência, competência, conformidade legal, certificação, conformidade com norma nem opinião de auditoria."

    doc.save(out_path)
    return count


v1.add_footer = add_footer
v1.core.add_footer = add_footer
v1.build_docx = build_docx
v1.core.build_docx = build_docx

if __name__ == "__main__":
    raise SystemExit(v1.core.main())
