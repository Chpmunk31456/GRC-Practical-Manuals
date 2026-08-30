#!/usr/bin/env python3
"""Generate trilingual Manual 13 SOX ITGC / ICFR publication-QA candidates."""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import generate_ccpa_cpra_manual_12_publication as base

ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "04-regulatory-compliance" / "SOX_ITGC_ICFR_Controlled_Implementation"
BASELINE = ROOT / ".compliance" / "sox-itgc-icfr-manual-13-baseline.json"
core = base.core

LANG_META = {
    "en": {
        "title": "Manual 13 - SOX ITGC / ICFR Controlled Implementation",
        "filename": "Manual_13_SOX_ITGC_ICFR_Controlled_Implementation_EN",
        "status": "CONTROLLED PUBLICATION QA CANDIDATE",
        "implementation_heading": "Implementation paths",
        "chapter_heading": "Controlled 32-chapter manual",
        "control": "Controlled source revision: {source_head} | Language: en | SOX/SEC/PCAOB authority boundaries retained.",
        "boundary": "Legal and audit boundary: This manual provides practical ICFR technology-control implementation guidance. It does not provide organization-specific legal advice, determine materiality, filer status, exemptions, deficiency classification, auditor scope, or guarantee SOX compliance. Current authoritative sources and competent human judgment remain controlling.",
        "figure_title": "Manual 13 memory graphic",
        "figure_caption": "Figure {number}. SOX ITGC / ICFR implementation memory graphic",
    },
    "es-419": {
        "title": "Manual 13 - Implementación Controlada SOX ITGC / ICFR",
        "filename": "Manual_13_SOX_ITGC_ICFR_Controlled_Implementation_ES-419",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACION",
        "implementation_heading": "Rutas de implementación",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "control": "Revisión de fuente controlada: {source_head} | Idioma: es-419 | se mantienen los límites de autoridad SOX/SEC/PCAOB.",
        "boundary": "Límite jurídico y de auditoría: Este manual ofrece guía práctica para controles tecnológicos de ICFR. No proporciona asesoría jurídica específica, no determina materialidad, condición de emisor, exenciones, clasificación de deficiencias, alcance del auditor ni garantiza cumplimiento SOX. Las fuentes vigentes y el juicio humano competente siguen siendo controlantes.",
        "figure_title": "Gráfico de memoria del Manual 13",
        "figure_caption": "Figura {number}. Gráfico de memoria de implementación SOX ITGC / ICFR",
    },
    "pt-BR": {
        "title": "Manual 13 - Implementação Controlada SOX ITGC / ICFR",
        "filename": "Manual_13_SOX_ITGC_ICFR_Controlled_Implementation_PT-BR",
        "status": "CANDIDATO CONTROLADO PARA QA DE PUBLICACAO",
        "implementation_heading": "Caminhos de implementação",
        "chapter_heading": "Manual controlado de 32 capítulos",
        "control": "Revisão da fonte controlada: {source_head} | Idioma: pt-BR | limites de autoridade SOX/SEC/PCAOB mantidos.",
        "boundary": "Limite jurídico e de auditoria: Este manual oferece orientação prática para controles tecnológicos de ICFR. Não fornece aconselhamento jurídico específico, não determina materialidade, status do emissor, isenções, classificação de deficiências, escopo do auditor nem garante conformidade SOX. Fontes vigentes e julgamento humano competente permanecem controlantes.",
        "figure_title": "Gráfico de memória do Manual 13",
        "figure_caption": "Figura {number}. Gráfico de memória de implementação SOX ITGC / ICFR",
    },
}


def source_dir(language: str) -> Path:
    folder = "English" if language == "en" else ("Spanish_es-419" if language == "es-419" else "Portuguese_pt-BR")
    return MANUAL / folder / "source"


def find_localized_chapters(language: str):
    chapters: dict[int, str] = {}
    used: list[str] = []
    skip = {"RUTAS_DE_IMPLEMENTACION_MANUAL_13.md", "CAMINHOS_DE_IMPLEMENTACAO_MANUAL_13.md"}
    for path in sorted(source_dir(language).glob("*.md")):
        if path.name in skip:
            continue
        found = core.split_chapters(path.read_text(encoding="utf-8"))
        if not found:
            continue
        for number, body in found.items():
            if number in chapters and chapters[number] != body:
                raise ValueError(f"conflicting chapter {number} for {language}: {path}")
            chapters[number] = body
        used.append(str(path.relative_to(ROOT)))
    expected = set(range(1, 33))
    if set(chapters) != expected:
        raise ValueError(f"{language} chapter inventory invalid: {sorted(chapters)}")
    return "\n".join(chapters[n].rstrip() for n in range(1, 33)) + "\n", used


def find_implementation(language: str):
    if language == "en":
        path = MANUAL / "MANUAL_13_IMPLEMENTATION_PATHS.md"
    elif language == "es-419":
        path = source_dir(language) / "RUTAS_DE_IMPLEMENTACION_MANUAL_13.md"
    else:
        path = source_dir(language) / "CAMINHOS_DE_IMPLEMENTACAO_MANUAL_13.md"
    text = path.read_text(encoding="utf-8")
    if len(core.MERMAID_RE.findall(text)) != 3:
        raise ValueError(f"expected exactly three memory graphics for {language}")
    markers = text.count("**Accessible explanation:**") + text.count("**Explicación accesible:**") + text.count("**Explicação acessível:**")
    if markers != 3:
        raise ValueError(f"expected three accessible explanations for {language}")
    return text, str(path.relative_to(ROOT))


def add_footer(section, language: str):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Manual 13 | {LANG_META[language]['status']} | ")
    core.set_run_font(run, language, size=8)
    field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE"); p._p.append(field)


def build_docx(source: core.EditionSource, out_path: Path, image_dir: Path, source_head: str):
    language = source.language
    meta = LANG_META[language]
    doc = Document()
    core.set_document_defaults(doc, language)
    add_footer(doc.sections[0], language)
    props = doc.core_properties
    props.title = meta["title"]
    props.subject = f"Manual 13 controlled publication-QA artifact; language {language}; source head {source_head}"
    props.author = "Alberto (Al) Leiva"
    props.keywords = "SOX, ICFR, ITGC, Section 302, Section 404, SEC, PCAOB, evidence, access, change management, computer operations"
    props.comments = "Automated generation does not provide legal advice, an audit opinion, deficiency classification, or a guarantee of SOX compliance."

    title = doc.add_paragraph(style="Title"); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["title"]); core.set_run_font(run, language, size=22, bold=True)
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta["status"]); core.set_run_font(run, language, size=10, bold=True)
    control = doc.add_paragraph(); control.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = control.add_run(meta["control"].format(source_head=source_head)); core.set_run_font(run, language, size=8.5)
    boundary = doc.add_paragraph(); run = boundary.add_run(meta["boundary"]); core.set_run_font(run, language, size=9, bold=True)

    graphic_counter = [0]
    h = doc.add_paragraph(meta["implementation_heading"], style="Heading 1"); core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.implementation_text, language, image_dir, graphic_counter, "Implementation")
    doc.add_section(WD_SECTION.NEW_PAGE); add_footer(doc.sections[-1], language)
    h = doc.add_paragraph(meta["chapter_heading"], style="Heading 1"); core.set_paragraph_keep(h, keep_next=True)
    core.add_markdown(doc, source.chapter_text, language, image_dir, graphic_counter, "Chapter")
    if graphic_counter[0] != 3:
        raise ValueError(f"{language} expected exactly three memory graphics")
    for p in doc.paragraphs:
        for r in p.runs:
            core.set_run_font(r, language)
    doc.save(out_path)
    return graphic_counter[0]


def write_reports(output_root: Path, sources: dict[str, core.EditionSource], file_results: dict, page_rows: list[dict], source_head: str):
    qa_dir = output_root / "qa"; qa_dir.mkdir(parents=True, exist_ok=True)
    manifest = {
        "schema_version": "1.0",
        "manual": "Manual 13 - SOX ITGC / ICFR Controlled Implementation",
        "source_head": source_head,
        "source_branch": "build/manual13-sox-icfr-2026",
        "assurance_boundary": "Automated QA supports integrity and review but does not provide legal advice, an audit opinion, materiality/deficiency conclusions, or replace competent review.",
        "editions": {},
    }
    for language, source in sources.items():
        manifest["editions"][language] = {
            "chapter_sources": source.chapter_files,
            "implementation_source": source.implementation_file,
            "chapter_count": 32,
            "mermaid_source_blocks": len(core.MERMAID_RE.findall(source.implementation_text)),
            "artifacts": file_results[language],
        }
    (qa_dir / "MANUAL_13_PUBLICATION_REPORT.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    (qa_dir / "MANUAL_13_PUBLICATION_REPORT.md").write_text(
        "# Manual 13 Publication Processing Report\n\n"
        f"- Controlled source head: `{source_head}`\n"
        "- Candidate languages: en, es-419, pt-BR\n"
        "- Automated document-processing status: **PASS**\n"
        "- Release status: **QA CANDIDATE**\n\n"
        "Automated checks confirm 32 chapters per edition, three implementation graphics with accessible explanations, DOCX/PDF processing, page extraction/rendering, and SHA-256 provenance. Genuine-human review requirements retained by repository control remain separate.\n",
        encoding="utf-8",
    )
    with (qa_dir / "MANUAL_13_PAGE_QA.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["pdf", "page", "width_pt", "height_pt", "text_chars", "render", "automated_status"])
        writer.writeheader(); writer.writerows(page_rows)
    candidates = sorted((output_root / "publication").glob("*.docx")) + sorted((output_root / "publication").glob("*.pdf"))
    with (qa_dir / "MANUAL_13_SHA256SUMS.txt").open("w", encoding="utf-8") as handle:
        for path in candidates:
            handle.write(f"{core.sha256(path)}  {path.name}\n")


core.MANUAL = MANUAL
core.BASELINE = BASELINE
core.LANG_META = LANG_META
core.find_localized_chapters = find_localized_chapters
core.find_implementation = find_implementation
core.add_footer = add_footer
core.build_docx = build_docx
core.write_reports = write_reports

if __name__ == "__main__":
    raise SystemExit(core.main())
