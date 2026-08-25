#!/usr/bin/env python3
"""Fail-closed document and publication QA for Manual 02 candidates."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import zipfile
from collections import Counter
from pathlib import Path

import fitz
from docx import Document
from lxml import etree
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "02-management-systems" / "ISO_IEC_42001_AIMS"
BASELINE = ROOT / ".compliance" / "iso-42001-manual-02-baseline.json"
OUTPUT_DIR = MANUAL / "publication" / "qa-candidate"
QA_DIR = ROOT / "qa" / "manual02-document-processing"

FILENAMES = {
    "en": "Manual_02_ISO_IEC_42001_AI_Management_System_EN",
    "es-419": "Manual_02_ISO_IEC_42001_AI_Management_System_ES-419",
    "pt-BR": "Manual_02_ISO_IEC_42001_AI_Management_System_PT-BR",
}
TITLES = {
    "en": "Manual 02 — ISO/IEC 42001 AI Management System Implementation",
    "es-419": "Manual 02 — Implementación del Sistema de Gestión de IA ISO/IEC 42001",
    "pt-BR": "Manual 02 — Implementação do Sistema de Gestão de IA ISO/IEC 42001",
}
STATUS = {
    "en": "PUBLICATION QA CANDIDATE — CONTROLLED ENGLISH SOURCE",
    "es-419": "DRAFT — HUMAN SEMANTIC REVIEW REQUIRED",
    "pt-BR": "DRAFT — HUMAN SEMANTIC REVIEW REQUIRED",
}
MASTER_PATHS = {
    "en": MANUAL / "English" / "ISO_IEC_42001_2023_Practical_AIMS_Manual_English_v1.0.md",
    "es-419": MANUAL / "translations" / "es-419" / "source" / f"{FILENAMES['es-419']}.md",
    "pt-BR": MANUAL / "translations" / "pt-BR" / "source" / f"{FILENAMES['pt-BR']}.md",
}

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
}


def rel(path: Path) -> str:
    return str(path.relative_to(ROOT))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def add_finding(
    findings: list[dict],
    language: str,
    issue: str,
    severity: str,
    proposed: str,
    page: int | None = None,
    chapter: str | None = None,
    gate: str | None = None,
) -> None:
    findings.append(
        {
            "language": language,
            "page": page,
            "chapter": chapter or "Preliminaries",
            "issue": issue,
            "severity": severity,
            "proposed_correction": proposed,
            "gate": gate,
        }
    )


def markdown_stats(language: str, path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    chapters = [
        int(value)
        for value in re.findall(r"(?m)^# (?:Chapter |Capítulo )?([1-9]|[12][0-9]|3[0-2])\. ", text)
    ]
    if language == "en":
        chapters = [
            int(value)
            for value in re.findall(r"(?m)^# ([1-9]|[12][0-9]|3[0-2])\. ", text)
        ]
    image_refs = re.findall(
        r'(?:<img\s+src="([^"]+)"[^>]*\salt="([^"]+)"\s*/?>|!\[([^\]]+)\]\(([^)]+)\))',
        text,
    )
    images = []
    for html_src, html_alt, md_alt, md_src in image_refs:
        images.append({"source": html_src or md_src, "alt": html_alt or md_alt})
    tables = len(re.findall(r"(?m)^\|.*\|$", text))
    headings = Counter(len(match.group(1)) for match in re.finditer(r"(?m)^(#{1,6})\s+", text))
    return {
        "path": rel(path),
        "chapters": chapters,
        "chapter_count": len(chapters),
        "images": images,
        "image_count": len(images),
        "table_rows": tables,
        "headings": dict(sorted(headings.items())),
        "text": text,
    }


def docx_stats(language: str, path: Path, findings: list[dict]) -> dict:
    document = Document(path)
    with zipfile.ZipFile(path) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        styles_xml = etree.fromstring(archive.read("word/styles.xml"))
        core_xml = etree.fromstring(archive.read("docProps/core.xml"))
        header_documents = [
            etree.fromstring(archive.read(name))
            for name in archive.namelist()
            if re.fullmatch(r"word/header\d+\.xml", name)
        ]
        footer_documents = [
            etree.fromstring(archive.read(name))
            for name in archive.namelist()
            if re.fullmatch(r"word/footer\d+\.xml", name)
        ]
        media = [name for name in archive.namelist() if name.startswith("word/media/")]

    tables = document_xml.xpath(".//w:tbl", namespaces=NS)
    one_cell = 0
    header_rows = 0
    merged_cells = len(document_xml.xpath(".//w:gridSpan | .//w:vMerge", namespaces=NS))
    for table in tables:
        rows = table.xpath("./w:tr", namespaces=NS)
        first_cells = rows[0].xpath("./w:tc", namespaces=NS) if rows else []
        if len(rows) == 1 and len(first_cells) == 1:
            one_cell += 1
        if rows and rows[0].xpath("./w:trPr/w:tblHeader", namespaces=NS):
            header_rows += 1

    images = document_xml.xpath(".//wp:docPr", namespaces=NS)
    empty_alt = [item.get("name") for item in images if not (item.get("descr") or "").strip()]
    textboxes = document_xml.xpath(".//w:txbxContent", namespaces=NS)
    hyperlinks = document_xml.xpath(".//w:hyperlink", namespaces=NS)
    hyperlink_texts = ["".join(node.xpath(".//w:t/text()", namespaces=NS)).strip() for node in hyperlinks]
    meaningless_links = [value for value in hyperlink_texts if not value or value.casefold() in {"here", "click here", "aquí", "aqui"}]
    title_nodes = core_xml.xpath("./dc:title/text()", namespaces=NS)
    core_title = title_nodes[0] if title_nodes else ""
    lang_values = styles_xml.xpath(".//w:lang/@w:val", namespaces=NS)
    instr = " ".join(document_xml.xpath(".//w:instrText/text()", namespaces=NS))
    footer_instr = " ".join(
        " ".join(document.xpath(".//w:instrText/text()", namespaces=NS))
        for document in footer_documents
    )
    header_text = " ".join(
        " ".join(document.xpath(".//w:t/text()", namespaces=NS))
        for document in header_documents
    )

    heading_levels = []
    heading_counts = Counter()
    list_count = 0
    for paragraph in document.paragraphs:
        style = paragraph.style.name or ""
        match = re.fullmatch(r"Heading ([1-9])", style)
        if match:
            level = int(match.group(1))
            heading_levels.append(level)
            heading_counts[level] += 1
        if style.startswith("List") or paragraph._p.xpath("./w:pPr/w:numPr"):
            list_count += 1

    critical_checks = {
        "title": core_title == TITLES[language],
        "tables": len(tables) == 33,
        "repeat_headers": header_rows == len(tables),
        "layout_tables": one_cell == 0,
        "images": len(images) == 10 and len(media) >= 10,
        "alt_text": not empty_alt,
        "headings": heading_counts[1] > 0 and heading_counts[2] > 0 and not any(
            next_level > current_level + 1
            for current_level, next_level in zip(heading_levels, heading_levels[1:])
        ),
        "lists": list_count > 0,
        "textboxes": len(textboxes) == 0,
        "links": len(hyperlinks) > 0 and not meaningless_links,
        "language": bool(lang_values),
        "toc": "TOC" in instr,
        "page_fields": "PAGE" in footer_instr and "NUMPAGES" in footer_instr,
        "header_status": STATUS[language] in header_text,
    }
    for check, passed in critical_checks.items():
        if not passed:
            add_finding(
                findings,
                language,
                f"DOCX accessibility/control check failed: {check}",
                "High",
                "Correct the DOCX generation/post-processing control and regenerate.",
                gate="Gate 2",
            )
    if merged_cells:
        add_finding(
            findings,
            language,
            f"DOCX contains {merged_cells} merged-cell markers; verify they do not impair navigation.",
            "Medium",
            "Review the affected tables with a screen reader and split cells if meaning is harmed.",
            gate="Gate 2",
        )

    return {
        "path": rel(path),
        "size_bytes": path.stat().st_size,
        "paragraphs": len(document.paragraphs),
        "heading_counts": dict(sorted(heading_counts.items())),
        "list_paragraphs": list_count,
        "tables": len(tables),
        "repeat_header_rows": header_rows,
        "layout_tables": one_cell,
        "merged_cell_markers": merged_cells,
        "images": len(images),
        "media_files": len(media),
        "svg_media_files": sum(name.endswith(".svg") for name in media),
        "empty_alt_text": len(empty_alt),
        "textboxes": len(textboxes),
        "hyperlinks": len(hyperlinks),
        "core_title": core_title,
        "language_values": sorted(set(lang_values)),
        "checks": critical_checks,
    }


def embedded_font_summary(reader: PdfReader) -> tuple[int, int]:
    total = 0
    embedded = 0
    seen: set[tuple[int, int]] = set()
    for page in reader.pages:
        resources = page.get("/Resources")
        if not resources:
            continue
        resources = resources.get_object()
        fonts = resources.get("/Font")
        if not fonts:
            continue
        for ref in fonts.get_object().values():
            key = (getattr(ref, "idnum", id(ref)), getattr(ref, "generation", 0))
            if key in seen:
                continue
            seen.add(key)
            total += 1
            font = ref.get_object()
            descriptor = font.get("/FontDescriptor")
            if descriptor:
                descriptor = descriptor.get_object()
                if any(name in descriptor for name in ("/FontFile", "/FontFile2", "/FontFile3")):
                    embedded += 1
    return total, embedded


def chapter_for_page(text: str, current: str) -> str:
    matches = re.findall(r"(?m)^(?:Chapter |Capítulo )?([1-9]|[12][0-9]|3[0-2])\.\s+[^\n]+", text)
    if matches:
        return f"Chapter {matches[-1]}"
    return current


def render_contact_sheet(document: fitz.Document, language: str, output_dir: Path) -> Path:
    thumbs: list[Image.Image] = []
    for index, page in enumerate(document):
        pix = page.get_pixmap(matrix=fitz.Matrix(0.34, 0.34), alpha=False)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        canvas = Image.new("RGB", (image.width, image.height + 22), "white")
        canvas.paste(image, (0, 22))
        draw = ImageDraw.Draw(canvas)
        draw.text((6, 4), f"{language} · page {index + 1}", fill="black", font=ImageFont.load_default())
        thumbs.append(canvas)
    columns = 6
    rows = (len(thumbs) + columns - 1) // columns
    cell_w = max(image.width for image in thumbs)
    cell_h = max(image.height for image in thumbs)
    sheet = Image.new("RGB", (columns * cell_w, rows * cell_h), "#d8dee4")
    for index, image in enumerate(thumbs):
        sheet.paste(image, ((index % columns) * cell_w, (index // columns) * cell_h))
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"manual02_{language}_contact_sheet.jpg"
    sheet.save(path, quality=88, optimize=True)
    return path


def pdf_stats(
    language: str,
    path: Path,
    findings: list[dict],
    page_records: list[dict],
    contact_dir: Path,
) -> dict:
    reader = PdfReader(path)
    if reader.is_encrypted:
        add_finding(findings, language, "PDF is unexpectedly encrypted.", "High", "Export without encryption.", gate="Gate 3")
    title = reader.metadata.title if reader.metadata else ""
    if title != TITLES[language]:
        add_finding(findings, language, "PDF title metadata is incorrect.", "High", "Set the DOCX title property and reconvert.", gate="Gate 3")
    try:
        outline_count = len(reader.outline or [])
    except Exception:
        outline_count = 0
    if outline_count == 0:
        add_finding(findings, language, "PDF has no navigation outline/bookmarks.", "Medium", "Enable heading bookmark export.", gate="Gate 3")
    font_total, font_embedded = embedded_font_summary(reader)
    if font_total and font_embedded == 0:
        add_finding(findings, language, "No embedded PDF fonts detected.", "High", "Embed or safely substitute fonts during PDF export.", gate="Gate 3")

    document = fitz.open(path)
    sizes = Counter((round(page.rect.width, 1), round(page.rect.height, 1)) for page in document)
    expected_size = sizes.most_common(1)[0][0]
    current_chapter = "Preliminaries"
    external_links = 0
    total_images = 0
    selectable_pages = 0
    page_failures = 0

    for index, page in enumerate(document, start=1):
        text = page.get_text("text")
        blocks = page.get_text("blocks")
        images = page.get_images(full=True)
        drawings = page.get_drawings()
        total_images += len(images)
        links = page.get_links()
        external_links += sum(1 for link in links if str(link.get("uri", "")).startswith(("http://", "https://")))
        if text.strip():
            selectable_pages += 1
        current_chapter = chapter_for_page(text, current_chapter)

        issues: list[tuple[str, str, str]] = []
        if len(text.strip()) < 20:
            issues.append(("Blank or non-searchable page", "High", "Remove the blank page or restore selectable text."))
        if any(marker in text for marker in ("�", "\ufffd", "(cid:")):
            issues.append(("Missing/corrupted glyph marker", "High", "Use a Unicode-safe embedded font and reconvert."))
        if (round(page.rect.width, 1), round(page.rect.height, 1)) != expected_size:
            issues.append(("Unexpected page size/orientation change", "High", "Normalize the section page setup and reconvert."))

        for block in blocks:
            x0, y0, x1, y1 = block[:4]
            if x0 < -1 or y0 < -1 or x1 > page.rect.width + 1 or y1 > page.rect.height + 1:
                issues.append(("Text block extends beyond the page boundary", "High", "Correct margins or table wrapping."))
                break

        for image in images:
            xref = image[0]
            for rect in page.get_image_rects(xref):
                if rect.x0 < -1 or rect.y0 < -1 or rect.x1 > page.rect.width + 1 or rect.y1 > page.rect.height + 1:
                    issues.append(("Graphic extends beyond the page boundary", "High", "Resize the graphic without changing its aspect ratio."))
                if rect.width < 250 or rect.height < 120:
                    issues.append(("Informative graphic may be unreadably small", "Medium", "Increase the rendered figure size while preserving aspect ratio."))

        caption_present = bool(re.search(r"(?m)^(?:Figure|Figura)\s+\d+\.", text))
        if caption_present and not images and not drawings:
            issues.append(("Figure caption is separated from its graphic", "High", "Keep the caption with the corresponding image."))

        pix = page.get_pixmap(matrix=fitz.Matrix(0.45, 0.45), colorspace=fitz.csGRAY, alpha=False)
        sample = pix.samples
        nonwhite = sum(1 for value in sample if value < 245) / max(1, len(sample))
        if nonwhite < 0.002:
            issues.append(("Page raster is effectively blank", "High", "Remove the unintended page or restore its content."))
        if nonwhite > 0.72:
            issues.append(("Page raster is unusually dense; overlap possible", "Medium", "Inspect for overlapping text or graphics."))

        status = "PASS" if not issues else "FAIL"
        if issues:
            page_failures += 1
            for issue, severity, correction in issues:
                add_finding(findings, language, issue, severity, correction, index, current_chapter, "Gate 4")
        page_records.append(
            {
                "language": language,
                "page": index,
                "chapter": current_chapter,
                "status": status,
                "text_characters": len(text.strip()),
                "images": len(images),
                "vector_drawings": len(drawings),
                "external_links": sum(1 for link in links if str(link.get("uri", "")).startswith(("http://", "https://"))),
                "nonwhite_ratio": round(nonwhite, 5),
                "issues": "; ".join(issue for issue, _, _ in issues),
            }
        )

    if selectable_pages != len(document):
        add_finding(findings, language, "Not every PDF page contains selectable text.", "High", "Restore the text layer and reconvert.", gate="Gate 3")
    if external_links == 0:
        add_finding(findings, language, "No preserved external PDF links detected.", "High", "Preserve hyperlinks during conversion.", gate="Gate 3")

    docx_document = Document(path.with_suffix(".docx"))
    header_pairs = []
    for table in docx_document.tables:
        if table.rows and len(table.rows[0].cells) >= 2:
            left = table.rows[0].cells[0].text.strip()
            right = table.rows[0].cells[1].text.strip()
            if len(left) >= 3 and len(right) >= 3:
                header_pairs.append((left, right))
    aligned_headers = 0
    for left, right in header_pairs:
        aligned = False
        for page in document:
            left_rects = page.search_for(left)
            right_rects = page.search_for(right)
            if any(abs(a.y0 - b.y0) <= 8 for a in left_rects for b in right_rects):
                aligned = True
                break
        aligned_headers += int(aligned)
    header_alignment_ratio = aligned_headers / max(1, len(header_pairs))
    if header_alignment_ratio < 0.8:
        add_finding(
            findings,
            language,
            f"Only {aligned_headers}/{len(header_pairs)} data-table header pairs render on the same row.",
            "High",
            "Repair table geometry so cell text remains inside the visible table grid, then reconvert.",
            gate="Gate 4",
        )

    contact_sheet = render_contact_sheet(document, language, contact_dir)
    return {
        "path": rel(path),
        "size_bytes": path.stat().st_size,
        "pages": len(document),
        "selectable_text_pages": selectable_pages,
        "page_size_counts": {f"{width}x{height}": count for (width, height), count in sizes.items()},
        "outline_entries": outline_count,
        "external_links": external_links,
        "placed_images": total_images,
        "fonts": font_total,
        "embedded_fonts": font_embedded,
        "metadata_title": title,
        "page_failures": page_failures,
        "aligned_table_headers": aligned_headers,
        "table_header_pairs": len(header_pairs),
        "table_header_alignment_ratio": round(header_alignment_ratio, 4),
        "contact_sheet": str(contact_sheet),
    }


def parity_findings(stats: dict, findings: list[dict]) -> dict:
    expected = list(range(1, 33))
    checks: dict[str, bool] = {}
    for language in ("en", "es-419", "pt-BR"):
        md = stats[language]["markdown"]
        docx = stats[language]["docx"]
        checks[f"{language}_chapters"] = md["chapters"] == expected
        checks[f"{language}_figures"] = md["image_count"] == 10 and docx["images"] == 10
        checks[f"{language}_tables"] = docx["tables"] == 33
        if not checks[f"{language}_chapters"]:
            add_finding(findings, language, "Chapter sequence is not exactly 1–32.", "High", "Restore the controlled source order.", gate="Gate 5")
        if not checks[f"{language}_figures"]:
            add_finding(findings, language, "Figure count is not 10 across master and DOCX.", "High", "Restore the missing or extra figure.", gate="Gate 5")
        if not checks[f"{language}_tables"]:
            add_finding(findings, language, "DOCX data-table count is not 33.", "High", "Reconcile tables with the English controlled source.", gate="Gate 5")

    localized_terms = {
        "es-419": ["Declaración de Aplicabilidad", "evaluación de riesgos de IA", "evaluación de impacto de sistemas de IA", "auditoría interna", "evidencia"],
        "pt-BR": ["Declaração de Aplicabilidade", "avaliação de riscos de IA", "avaliação de impacto de sistemas de IA", "auditoria interna", "evidência"],
    }
    for language, terms in localized_terms.items():
        text = stats[language]["markdown"]["text"]
        checks[f"{language}_controlled_terms"] = all(term.casefold() in text.casefold() for term in terms)
        checks[f"{language}_draft_status"] = STATUS[language] in text
        expected_path = f"assets/{language}/media/"
        sources = [item["source"] for item in stats[language]["markdown"]["images"]]
        checks[f"{language}_own_graphics"] = all(expected_path in source for source in sources) and not any("assets/English/" in source for source in sources)
        checks[f"{language}_risk_impact_distinct"] = terms[1].casefold() in text.casefold() and terms[2].casefold() in text.casefold()
        for check in ("controlled_terms", "draft_status", "own_graphics", "risk_impact_distinct"):
            if not checks[f"{language}_{check}"]:
                add_finding(findings, language, f"Trilingual parity check failed: {check}", "High", "Reconcile mechanically with the controlled source; send substantive language questions to human review.", gate="Gate 5")

    for source_id in ("ISO/IEC 42001", "ISO/IEC 42005", "ISO/IEC 42006", "ISO/IEC 23894", "ISO 19011"):
        checks[f"reference_{source_id}"] = all(source_id in stats[language]["markdown"]["text"] for language in stats)
        if not checks[f"reference_{source_id}"]:
            add_finding(findings, "trilingual", f"Source reference disappeared in one or more editions: {source_id}", "High", "Restore the controlled source reference.", gate="Gate 5")
    return checks


def gate_status(findings: list[dict], gate: str) -> str:
    blocking = [item for item in findings if item.get("gate") == gate and item["severity"] in {"Critical", "High"}]
    return "FAIL" if blocking else "PASS"


def write_reports(
    stats: dict,
    findings: list[dict],
    page_records: list[dict],
    gates: dict,
    source_commit: str,
    source_branch: str,
    generation_date: str,
) -> None:
    QA_DIR.mkdir(parents=True, exist_ok=True)
    json_path = QA_DIR / "ISO_IEC_42001_MANUAL_02_DOCUMENT_PROCESSING_REPORT.json"
    csv_path = QA_DIR / "ISO_IEC_42001_MANUAL_02_PAGE_QA.csv"
    md_path = QA_DIR / "ISO_IEC_42001_MANUAL_02_DOCUMENT_PROCESSING_REPORT.md"
    checksums_path = QA_DIR / "ISO_IEC_42001_MANUAL_02_SHA256SUMS.txt"

    serializable_stats = {}
    for language, language_stats in stats.items():
        if language == "parity":
            serializable_stats[language] = language_stats
            continue
        serializable_stats[language] = {
            **language_stats,
            "markdown": {
                key: value for key, value in language_stats["markdown"].items() if key != "text"
            },
        }

    payload = {
        "manual": "Manual 02 — ISO/IEC 42001 AI Management System Implementation",
        "source_branch": source_branch,
        "source_commit": source_commit,
        "generation_date": generation_date,
        "release_status": "blocked-by-translation-review",
        "gates": gates,
        "languages": serializable_stats,
        "findings": findings,
        "page_records": page_records,
        "human_review": {
            "status": "OPEN",
            "languages": ["es-419", "pt-BR"],
            "automatic_closure_permitted": False,
        },
    }
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(page_records[0]))
        writer.writeheader()
        writer.writerows(page_records)

    artifacts = [
        OUTPUT_DIR / f"{FILENAMES[language]}.{extension}"
        for language in ("en", "es-419", "pt-BR")
        for extension in ("docx", "pdf")
    ] + [MASTER_PATHS["es-419"], MASTER_PATHS["pt-BR"]]
    checksums_path.write_text(
        "".join(f"{sha256(path)}  {rel(path)}\n" for path in artifacts), encoding="utf-8"
    )

    visual_failures = [item for item in findings if item.get("gate") == "Gate 4"]
    accessibility_failures = [item for item in findings if item.get("gate") == "Gate 2"]
    parity_failures = [item for item in findings if item.get("gate") == "Gate 5"]
    lines = [
        "# DOCUMENT PROCESSING REPORT",
        "",
        f"- Source branch: `{source_branch}`",
        f"- Source commit: `{source_commit}`",
        f"- Generation date: `{generation_date}`",
        "- Localized status: **DRAFT — HUMAN SEMANTIC REVIEW REQUIRED**",
        "- Assurance boundary: generation and QA do not establish ISO authorization, certification, conformity, legal compliance, or audit assurance.",
        "",
        "## Publication artifacts",
        "",
        "| Language | Markdown master | DOCX | PDF | Pages |",
        "|---|---|---|---|---:|",
    ]
    for language in ("en", "es-419", "pt-BR"):
        lines.append(
            f"| {language} | `{stats[language]['markdown']['path']}` | `{stats[language]['docx']['path']}` | `{stats[language]['pdf']['path']}` | {stats[language]['pdf']['pages']} |"
        )
    lines += ["", "## Quality gates", "", "| Gate | Result |", "|---|---|"]
    lines.extend(f"| {gate} | {value} |" for gate, value in gates.items())
    lines += [
        "",
        "## Accessibility findings",
        "",
        "No unresolved high/critical DOCX accessibility finding." if not accessibility_failures else f"{len(accessibility_failures)} finding(s); see JSON report.",
        "",
        "## Visual QA findings",
        "",
        "Every generated PDF page has a recorded page-level result. No unresolved high/critical visual finding." if not visual_failures else f"{len(visual_failures)} page-level finding(s); see CSV and JSON reports.",
        "",
        "## Trilingual parity findings",
        "",
        "All editions contain chapters 1–32 in order, ten figures, 33 data tables, controlled references, and the required risk/impact and applicability terminology." if not parity_failures else f"{len(parity_failures)} parity finding(s); see JSON report.",
        "",
        "## Human-review dependency",
        "",
        "The Spanish and Brazilian Portuguese semantic/terminology review gate remains OPEN. These artifacts are layout and accessibility QA candidates only and are not release-ready.",
        "",
        "## QA workflows run",
        "",
        "- Controlled-source, human-gate, source-registry, and manual-catalog repository checks",
        "- DOCX semantic/accessibility package inspection",
        f"- PDF content, link, font, and page-by-page raster inspection ({len(page_records)} pages)",
        "- Trilingual chapter, figure, table, terminology, graphic-language, and reference parity",
        "- SHA-256 artifact manifest generation",
        "",
        "DOCUMENT PROCESSING STATUS: BLOCKED BY TRANSLATION REVIEW",
        "",
    ]
    md_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-commit", required=True)
    parser.add_argument("--source-branch", default="build/iso-iec-42001-manual-02-2026")
    parser.add_argument("--generation-date", required=True)
    parser.add_argument("--contact-dir", type=Path, default=Path("/tmp/manual02-contact-sheets"))
    args = parser.parse_args()

    baseline = json.loads(BASELINE.read_text(encoding="utf-8"))
    findings: list[dict] = []
    page_records: list[dict] = []
    stats: dict[str, dict] = {}

    for language in ("en", "es-419", "pt-BR"):
        master = MASTER_PATHS[language]
        docx = OUTPUT_DIR / f"{FILENAMES[language]}.docx"
        pdf = OUTPUT_DIR / f"{FILENAMES[language]}.pdf"
        for path in (master, docx, pdf):
            if not path.is_file() or path.stat().st_size == 0:
                add_finding(findings, language, f"Required artifact missing or empty: {rel(path)}", "Critical", "Regenerate the controlled artifact.", gate="Gate 1")
        stats[language] = {
            "markdown": markdown_stats(language, master),
            "docx": docx_stats(language, docx, findings),
            "pdf": pdf_stats(language, pdf, findings, page_records, args.contact_dir),
        }

    expected_chapters = list(range(1, int(baseline["required_localized_chapters"]) + 1))
    for language in ("en", "es-419", "pt-BR"):
        if stats[language]["markdown"]["chapters"] != expected_chapters:
            add_finding(findings, language, "Controlled source chapter order failed.", "Critical", "Restore chapters 1–32 exactly once and in order.", gate="Gate 1")

    parity = parity_findings(stats, findings)
    stats["parity"] = parity
    gates = {
        "Gate 1 — Source integrity": gate_status(findings, "Gate 1"),
        "Gate 2 — DOCX generation": gate_status(findings, "Gate 2"),
        "Gate 3 — PDF generation": gate_status(findings, "Gate 3"),
        "Gate 4 — Page-by-page visual QA": gate_status(findings, "Gate 4"),
        "Gate 5 — Trilingual parity": gate_status(findings, "Gate 5"),
        "Gate 6 — Human semantic approval": "OPEN — NOT AUTOMATICALLY CLOSED",
        "Gate 7 — Release package": "BLOCKED",
    }
    write_reports(
        stats,
        findings,
        page_records,
        gates,
        args.source_commit,
        args.source_branch,
        args.generation_date,
    )
    print(json.dumps({"gates": gates, "findings": findings, "pages": len(page_records)}, ensure_ascii=False, indent=2))
    return 1 if any(value == "FAIL" for value in gates.values()) else 0


if __name__ == "__main__":
    raise SystemExit(main())
